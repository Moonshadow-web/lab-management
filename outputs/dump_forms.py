import docx, openpyxl, os, sys

BASE = r"d:\workbuddyprojects\网页版-生免速查工具\data\uploads\docs"

def dump_docx(path):
    print("="*80)
    print("DOCX:", os.path.basename(path))
    print("="*80)
    d = docx.Document(path)
    print("--- PARAGRAPHS (non-empty) ---")
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if t:
            print(f"[{i}|{p.style.name}] {t}")
    print(f"\n--- TABLES: {len(d.tables)} ---")
    for ti, tbl in enumerate(d.tables):
        print(f"\n### TABLE {ti}  rows={len(tbl.rows)} cols={len(tbl.columns)}")
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.replace('\n','\\n').strip() for c in row.cells]
            print(f"R{ri}: " + " | ".join(cells))

def dump_xlsx(path):
    print("="*80)
    print("XLSX:", os.path.basename(path))
    print("="*80)
    wb = openpyxl.load_workbook(path, data_only=False)
    for ws in wb.worksheets:
        print(f"\n### SHEET: {ws.title}  dims={ws.dimensions}  max_row={ws.max_row} max_col={ws.max_column}")
        for r in range(1, min(ws.max_row, 60)+1):
            vals = []
            for c in range(1, min(ws.max_column, 30)+1):
                v = ws.cell(row=r, column=c).value
                if v is not None:
                    vals.append(f"{openpyxl.utils.get_column_letter(c)}{r}={repr(v)}")
            if vals:
                print(" | ".join(vals))

targets = [
    "BG-SM-CZ-021-定性室内比对结果记录及分析报告表.docx",
    "BG-SM-CZ-022-定量室内比对报告.docx",
    "BG-SM-CZ-024-定量室内比对结果记录分析表_DXI800分析仪_.xlsx",
    "BG-SM-CZ-025-定量室内比对结果分析表_生化分析仪_.xlsx",
    "BG-SM-CZ-026-定量室内比对结果记录分析表_凝血分析仪_.xlsx",
    "BG-SM-CZ-027-定量室内比对结果记录分析表_早孕系列_.xlsx",
    "BG-SM-CZ-071-定量室内比对结果记录分析表_血气分析仪_.xlsx",
]

for t in targets:
    p = os.path.join(BASE, t)
    if not os.path.exists(p):
        print("MISSING:", t); continue
    if t.endswith(".docx"):
        dump_docx(p)
    else:
        dump_xlsx(p)
