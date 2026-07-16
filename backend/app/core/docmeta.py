"""从 Word 文档（.docx）中提取文件头元数据。

生免组 SOP 类文档（通用SOP/项目SOP/仪器SOP）首页通常为一张 6×5 的
表头表格，包含：文件编号、版本号、修订号、发布/审核/批准/实施日期、
编写者/审核者/批准者，以及标题、编写单位等。

本模块把这些字段解析出来供文件管理模块展示，并在「修改后上传新版本」
时自动记录到版本历史，实现可追溯。

解析策略：
1. 仅支持 .docx（OpenXML）。老版 .doc 为二进制格式，python-docx 无法读取，
   返回 {}（前端显示“—”）。后续可接 LibreOffice 转 docx 再解析。
2. 取文档前 3 个表格的所有单元格 + 前 20 个段落，组成有序 token 序列。
3. 对每个 token：若包含某字段关键词，则优先取同行内「key：value」形式的值；
   否则取序列中下一个非关键词 token 作为值。
"""
from __future__ import annotations

import os
import re
from typing import Optional

# 字段 -> 关键词（按特异性从高到低排列，避免“审核日期”误匹配“审核者”）
KEY_PATTERNS = {
    "doc_number": ["文件编号", "文件编码", "文件代号"],
    "doc_version": ["版本号", "版本：", "版次", "版本:"],
    "revision": ["修订号", "修订：", "修订:"],
    "issued_date": ["发布日期", "编制日期", "签发日期", "发布："],
    "audit_date": ["审核日期", "审查日期", "复核日期", "审核："],
    "approve_date": ["批准日期", "审批日期", "批准："],
    "effective_date": ["实施日期", "生效日期", "执行日期", "启用日期", "实施："],
    "author": ["编写者", "编制人", "编写人", "起草人", "编制："],
    "reviewer": ["审核者", "审查者", "审核人", "复核人", "审核:"],
    "approver": ["批准者", "审批人", "批准人", "审批:"],
}

# 所有字段的全部关键词（用于判断某个 token 是否“本身是个 key”）
_ALL_KEYWORDS = [kw for kws in KEY_PATTERNS.values() for kw in kws]


def _is_key_token(text: str) -> bool:
    return any(kw in text for kw in _ALL_KEYWORDS)


def _extract_number_from_filename(filename: str) -> Optional[str]:
    """从文件名提取文件编号（如 SM-SOP-534、BG-SM-CZ-024）。"""
    if not filename:
        return None
    base = os.path.splitext(os.path.basename(filename))[0]
    # 匹配开头的大写字母/数字 + 连字符组合，遇到非编码字符停止
    m = re.match(r"^([A-Z0-9]+(?:-[A-Z0-9]+)*)", base)
    if m:
        return m.group(1)
    return None


def _inline_value(text: str) -> Optional[str]:
    """提取 'key：value' / 'key:value' 中冒号后的值；没有则返回 None。"""
    for sep in ("：", ":"):
        if sep in text:
            val = text.split(sep, 1)[1].strip()
            # 去掉可能的尾随标签（如“编写者：金子铮”）
            return val or None
    return None


def _clean(value: str) -> str:
    value = value.strip()
    # 去掉可能的“页码：第1页 共17页”这类误取（含“页”且像页码）
    if "第" in value and "页" in value:
        return ""
    return value


def _build_record_table_meta(file_path: str, title: str = "") -> dict:
    """记录表格（category='记录表格'）按命名规则生成元数据。

    规则（用户确认）：
    - 表格编号 = 文件名前缀，如 BG-KS-CZ-901 / BG-SM-GL-006 / SM-CZ-044
      （去掉 Word 临时文件标记 ~$ 及其落盘后的脏前缀 __ / _$ 等）
    - 版本号 01，修订号 0
    - 前缀含 KS（科室）：编写人 刘书理 / 审核人 张婵媛 / 批准人 王学晶
    - 前缀含 SM（生免组）：编写人 金子铮 / 审核人 杨静 / 批准人 王学晶
    - 发布/审核日期 2025-03-01；批准/实施日期 2025-04-01
    """
    # 记录表格编号始终从文件名（file_path basename）提取，不依赖 title——
    # 因为标题字段已去掉编号前缀，但存储文件名仍保留前缀。
    candidate = os.path.basename(file_path or "").strip()
    if not candidate:
        candidate = (title or "").strip()
    # 清理 Word 临时文件标记 ~$ 及其落盘后的脏前缀（__ / _$ / ~$- 等），
    # 并去掉开头所有非字母数字字符（避免 “-SM-GL-003” 这类无法匹配）
    candidate = re.sub(r"^[^A-Za-z0-9]+", "", candidate.replace("~$", ""))
    base = os.path.splitext(candidate)[0]
    m = re.match(r"^(BG-[A-Z]{2}-[A-Z]{2,}-\d+|[A-Z]{2}-[A-Z]{2,}-\d+)", base)
    if not m:
        return {}
    number = m.group(1)
    upper = number.upper()
    if "SM" in upper:
        author, reviewer = "金子铮", "杨静"
    else:
        author, reviewer = "刘书理", "张婵媛"
    approver = "王学晶"
    return {
        "doc_number": number,
        "doc_version": "01",
        "revision": "0",
        "author": author,
        "reviewer": reviewer,
        "approver": approver,
        "issued_date": "2025-03-01",
        "audit_date": "2025-03-01",
        "approve_date": "2025-04-01",
        "effective_date": "2025-04-01",
    }


def parse_doc_metadata(file_path: str, title: str = "", category: str = "") -> dict:
    """解析 Word 文档头表元数据，返回字段字典（仅含解析到的字段）。

    支持 .docx 表头解析；对旧版 .doc 无法读取内容，但会尝试从文件名
    提取文件编号，使项目 SOP 等旧文档也能展示编号。
    """
    result: dict[str, str] = {}

    # 0) 记录表格：按命名规则生成元数据（不解析文件正文，避免污染）
    if category == "记录表格":
        return _build_record_table_meta(file_path, title)

    # 1) .docx：读取前 3 个表格 + 前 20 段落的表头元数据
    if file_path and file_path.lower().endswith(".docx") and os.path.exists(file_path):
        try:
            import docx
            doc = docx.Document(file_path)
        except Exception:
            doc = None

        if doc is not None:
            tokens: list[str] = []
            for tbl in doc.tables[:3]:
                for row in tbl.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            tokens.append(t)
            for par in doc.paragraphs[:20]:
                t = par.text.strip()
                if t:
                    tokens.append(t)

            for i, tok in enumerate(tokens):
                for field, kws in KEY_PATTERNS.items():
                    if field in result:
                        continue
                    for kw in kws:
                        if kw in tok:
                            inline = _inline_value(tok)
                            if inline:
                                result[field] = _clean(inline)
                                break
                            nxt = tokens[i + 1] if i + 1 < len(tokens) else ""
                            if nxt and not _is_key_token(nxt):
                                cleaned = _clean(nxt)
                                if cleaned:
                                    result[field] = cleaned
                                    break
                    else:
                        continue
                    break

            # 记录原始头表（仅第一表）用于审计/调试
            if doc.tables:
                raw = "\n".join(
                    " | ".join(c.text.strip() for c in r.cells)
                    for r in doc.tables[0].rows
                )
                if raw:
                    result["_raw_table"] = raw

    # 2) 文件名兜底：当文件编号仍为空且文件为 SOP 常见格式时，从文件名/标题提取。
    #    为保持记录表格等不被污染，兜底仅对 SOP 类文档生效。
    SOP_CATEGORIES = {"通用SOP", "项目SOP", "仪器SOP"}
    if (
        not result.get("doc_number")
        and file_path.lower().endswith((".docx", ".doc", ".pdf"))
        and category in SOP_CATEGORIES
    ):
        candidate = os.path.basename(file_path or "") or title or ""
        number = _extract_number_from_filename(candidate)
        if number:
            result["doc_number"] = number

    # 3) 强信号校验：文件编号或文件版本号至少有一个可信来源，才采信其余字段
    if "doc_number" not in result and "doc_version" not in result:
        return {}

    return result
