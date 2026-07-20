"""室内质控月小结（CZ-012）Word 报告生成：A4 横向，含 14 列表格 + 五段文字部分。

版式严格对应 CZ-012 表单：表格编号 BG-SM-CZ-012 置于页脚；年/月按「2026年06月」
（数字在前、文字在后、不用冒号）呈现；质控批号仅作为表格列，不单独成句。
"""
import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .comparison_report import _fill, _run_font, _add_footer_code, _heading

CZ012_HEADERS = [
    "项目", "质控批号", "单位", "水平", "靶值", "靶值SD", "靶值CV%",
    "均值", "SD", "CV%", "n", "失控数", "在控率", "质量目标（允许不精密度）",
]

# 各列宽度（cm）。横向 A4 可用宽 ≈ 29.7 − 1.8×2 = 26.1cm
_CZ012_WIDTHS = [3.0, 2.2, 1.4, 1.6, 2.0, 2.0, 2.0, 2.0, 2.0, 1.8, 1.2, 1.8, 1.8, 2.3]


def _fmt(v, nd=2):
    try:
        f = float(v)
    except (TypeError, ValueError):
        return v if v not in (None, "") else "—"
    return f"{f:.{nd}f}"


def build_docx(out_path, summaries, report, instrument_name, instrument_no, year, month):
    """生成 CZ-012 月小结 Word（A4 横向）并保存到 out_path，返回 out_path。

    summaries: 该 (仪器,年,月) 下的 QCMonthlySummary 列表
    report:    对应的 QCMonthlyReport（文字部分）；为 None 时文字段留空
    """
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    _add_footer_code(sec, "BG-SM-CZ-012", year)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run("生化免疫组室内质控月小结"), 15, bold=True)

    # 副标题：2026年06月　仪器：XXX（编号：YYY）
    sub = f"{year or '—'}年{str(month).zfill(2) if month else '—'}月　　仪器：{instrument_name or '—'}"
    if instrument_no:
        sub += f"（编号：{instrument_no}）"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(sub), 10.5)

    # 表单信息行：表格编号 / 科室 / 生效日期（与 CZ-012 表单一致）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(p.add_run(
        "表格编号：BG-SM-CZ-012    民航总医院检验科生化免疫组       生效日期：2026.08.01"
    ), 10.5)

    # 数据表格
    ncol = len(CZ012_HEADERS)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c, h in enumerate(CZ012_HEADERS):
        _fill(t.rows[0].cells[c], h, bold=True, size=9, align="center")
    for s in summaries:
        cells = t.add_row().cells
        vals = [
            s.test_item, s.lot_no, s.unit, s.level,
            _fmt(s.target_mean), _fmt(s.target_sd), f"{s.target_cv:.2f}%",
            _fmt(s.mean), _fmt(s.sd), f"{s.cv:.2f}%",
            s.n, s.out_of_control_count, f"{s.in_control_rate * 100:.1f}%",
            s.quality_goal,
        ]
        for c, v in enumerate(vals):
            _fill(cells[c], v, size=9)
    widths = [Cm(w) for w in _CZ012_WIDTHS]
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

    # 文字部分（无独立标题）
    sections = [
        ("一、仪器运行情况", report.operation_status if report else ""),
        ("二、各项目是否出现漂移或趋势性改变", report.drift_trend if report else ""),
        ("三、各项目CV%设置是否达标", report.cv_setting_ok if report else ""),
        ("四、各项目计算CV%是否达标", report.cv_calc_ok if report else ""),
        ("五、各项目质控频次是否达标", report.freq_ok if report else ""),
    ]
    for label, text in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        _run_font(p.add_run(label + "："), 11, bold=True)
        _run_font(p.add_run(text or "（未填写）"), 11)

    # 落款/签字审批区（左对齐，2026-07-19 调整：去除"签字"二字，改为普通落款）
    # 仪器日常管理人
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    _run_font(p.add_run("仪器日常管理人："), 11)

    # 质控总负责人审批意见（标签 + 空行）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    _run_font(p.add_run("质控总负责人审批意见："), 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(p.add_run(""), 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(p.add_run(""), 11)

    # 审批人 + 年月日（右下角落款，年月日向右缩进，左侧留空便于手填年份）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(8)
    _run_font(p.add_run("审批人："), 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run_font(p.add_run("年    月    日"), 11)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
