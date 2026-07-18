"""仪器间比对：结果计算 + 报告生成（docx，保留表格编号）+ HTML 预览。

计算逻辑：
- 定量：偏倚 = (val - ref)/ref × 100（relative）或 val - ref（absolute）；
  是否允许 = |偏倚| ≤ 允许TE。
- 定性：以参照仪器 5 例样本为基准，逐台比对仪器计算阴阳符合率。

报告版式严格参照对应 SOP 表单（BG-SM-CZ-021/022/024~027/071），保留表格编号。
"""

import json
import os
import re as _re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from ..models.comparison import ComparisonGroup, ComparisonPlan, ComparisonResult, ComparisonQualResult
from ..models.instrument import Instrument

REPORT_DIR = None  # 由 main 注入（DATA_DIR/comparison_reports）


# ---------------------------------------------------------------------------
# 工具
# ---------------------------------------------------------------------------
def _load_json(s, default):
    try:
        v = json.loads(s) if isinstance(s, str) else s
        return v if v is not None else default
    except Exception:
        return default


def disp_name(inst: Instrument) -> str:
    """人可识别的仪器显示名：优先规格型号(model)，其次名称。

    仪器档案里 name 多为通用名（如"全自动生化分析仪"），model 才是可识别的
    具体型号（如"贝克曼AU5821 A"、"TOP700A"）。比对模块统一用 model 展示。
    """
    if not inst:
        return ""
    model = (inst.model or "").strip()
    name = (inst.name or "").strip()
    if model and len(model) <= 40:
        return model
    return name or model or f"仪器{inst.id}"


def _instruments_of(db, group: ComparisonGroup):
    ids = _load_json(group.instrument_ids, [])
    ref_id = group.reference_instrument_id
    out = []
    seen = set()
    for iid in ids:
        if iid in seen:
            continue
        inst = db.get(Instrument, iid)
        if inst:
            seen.add(iid)
            out.append({"id": inst.id, "name": disp_name(inst), "model": inst.model or "",
                        "is_reference": inst.id == ref_id})
    if not any(x["is_reference"] for x in out) and ref_id and ref_id not in seen:
        inst = db.get(Instrument, ref_id)
        if inst:
            out.insert(0, {"id": inst.id, "name": disp_name(inst), "model": inst.model or "",
                           "is_reference": True})
    return out


def _compared_ids(group: ComparisonGroup):
    ids = _load_json(group.instrument_ids, [])
    return [i for i in ids if i != group.reference_instrument_id]


def _parse_float(x):
    try:
        return float(str(x).strip())
    except Exception:
        return None


def _compute_bias(ref, val, te, mode):
    """返回 (偏倚值或None, 是否允许或None)。"""
    rf = _parse_float(ref)
    vf = _parse_float(val)
    tf = _parse_float(te)
    if rf is None or vf is None or tf is None:
        return None, None
    if mode == "absolute":
        bias = vf - rf
    else:
        if rf == 0:
            return None, None
        bias = (vf - rf) / rf * 100.0
    accepted = abs(bias) <= tf + 1e-9
    return round(bias, 2), accepted


# ---------------------------------------------------------------------------
# 计算
# ---------------------------------------------------------------------------
def compute_data(db, group: ComparisonGroup, plan: ComparisonPlan):
    instruments = _instruments_of(db, group)
    ref_id = group.reference_instrument_id
    compared = _compared_ids(group)
    items = _load_json(group.items, [])

    if group.category == "定性":
        quals = {
            q.item: q
            for q in db.query(ComparisonQualResult).filter_by(plan_id=plan.id).all()
        }
        all_ids = [ins["id"] for ins in instruments]

        def _qual_applicable(it):
            ids = it.get("instrument_ids") or []
            if not ids:
                return set(all_ids)
            return {i for i in all_ids if i in set(ids)}

        matrix = []
        overall_min = 100.0
        # 仅纳入本次计划实际录入了定性结果的项目（比对分批次进行，未参与本次的不列）
        involved = set(quals.keys())
        for it in items:
            if it["name"] not in involved:
                continue
            qr = quals.get(it["name"])
            rj = _load_json(qr.results_json, {}) if qr else {}
            ref_res = rj.get(str(ref_id))
            app = _qual_applicable(it)
            insts = {}
            for ins in instruments:
                if ins["id"] not in app:
                    insts[ins["id"]] = {"name": ins["name"], "results": [], "agreement": None, "masked": True}
                    continue
                res = rj.get(str(ins["id"])) or []
                agreement = None
                if res and ref_res and len(res) == len(ref_res):
                    valid = sum(1 for a in res if a in ("P", "N"))
                    matches = sum(1 for a, b in zip(res, ref_res) if a == b and a in ("P", "N"))
                    agreement = round(matches / valid * 100, 1) if valid else None
                insts[ins["id"]] = {"name": ins["name"], "results": res, "agreement": agreement, "masked": False}
            matrix.append({"item": it["name"], "insts": insts})
        return {
            "category": "定性",
            "instruments": instruments,
            "compared": compared,
            "matrix": matrix,
            "ref_id": ref_id,
        }

    # 定量
    results = {
        (r.item, r.level): r
        for r in db.query(ComparisonResult).filter_by(plan_id=plan.id).all()
    }
    levels = group.levels or 5
    # 仅纳入本次计划实际录入了定量结果的项目（比对分批次进行，未参与本次的不列）
    involved = {name for (name, _lv) in results.keys()}
    all_inst_ids = [i["id"] for i in instruments]

    def _all_applicable(it):
        """该项目全部适用仪器 id（含参照仪器本身）；空 = 组内全部仪器。"""
        ids = it.get("instrument_ids") or []
        if not ids:
            return set(all_inst_ids)
        return {c for c in all_inst_ids if c in set(ids)}

    def _effective_ref(app):
        """找出该项目的实际参照仪器 id：优先组参照；若组参照不适用，取第一个可用的比对仪器。"""
        if ref_id in app:
            return ref_id
        for cid in compared:
            if cid in app:
                return cid
        return ref_id  # fallback（不应发生）

    def _compare_to_ids(app, eff_ref):
        """该项目实际参与计算的比对仪器 id（排除有效靶机自身）。"""
        return [cid for cid in compared if cid in app and cid != eff_ref]

    # 预计算每项目的有效靶机（供 summary / matrix / 前端使用）
    item_eff_refs = {}
    for it in items:
        if it["name"] not in involved:
            continue
        app = _all_applicable(it)
        eff_ref = _effective_ref(app)
        item_eff_refs[it["name"]] = eff_ref

    matrix = []
    for lv in range(1, levels + 1):
        rows = []
        for it in items:
            if it["name"] not in involved:
                continue
            r = results.get((it["name"], lv))
            ref_val = r.reference_value if r else ""
            eff_ref = item_eff_refs[it["name"]]
            app = _all_applicable(it)
            compare_ids = _compare_to_ids(app, eff_ref)
            insts = {}
            # 有效靶机（动态参照）：标记 is_ref=True，在报告中显示"参照"而非"/"
            v_ref = _load_json(r.values_json, {}).get(str(eff_ref), "") if r else ""
            insts[str(eff_ref)] = {"value": v_ref, "bias": None, "accepted": None, "masked": False, "is_ref": True}
            # 非靶机的比对仪器 → 正常计算偏倚
            for cid in compare_ids:
                v = _load_json(r.values_json, {}).get(str(cid), "") if r else ""
                bias, accepted = _compute_bias(ref_val, v, _resolve_te(it, lv), _resolve_mode(it, lv))
                insts[str(cid)] = {"value": v, "bias": bias, "accepted": accepted, "masked": False}
            # 其余不适用仪器 → masked（组参照仪不在 app 或其它不适用）
            for cid in [i for i in [ref_id] + compared if i != eff_ref and i not in compare_ids]:
                insts[str(cid)] = {"value": "", "bias": None, "accepted": None, "masked": True}
            rows.append({
                "item": it["name"],
                "label": it.get("label", ""),
                "te": _resolve_te(it, lv),
                "mode": _resolve_mode(it, lv),
                "ref": ref_val,
                "effective_ref_id": eff_ref,
                "insts": insts,
            })
        matrix.append({"level": lv, "rows": rows})

    # 汇总：分仪器——注意动态靶机下，某仪器若是某项目的实际参照，则该项目不归入该仪器的成绩
    summary = []
    for cid in compared:
        per_item = []
        for it in items:
            if it["name"] not in involved:
                continue
            eff_ref = item_eff_refs.get(it["name"])
            if cid == eff_ref:
                continue  # 该仪器此项是靶机（非比对仪器），不列入汇总
            app = _all_applicable(it)
            if cid not in app:
                continue
            level_info = []
            for lv in range(1, levels + 1):
                r = results.get((it["name"], lv))
                # 偏倚计算：对照该项目的实际靶机值（reference_value）
                v = _load_json(r.values_json, {}).get(str(cid), "") if r else ""
                ref_val = r.reference_value if r else ""
                _, accepted = _compute_bias(r.reference_value if r else "", v, _resolve_te(it, lv), _resolve_mode(it, lv))
                level_info.append({"ref": ref_val, "ok": accepted is True})
            per_item.append({"item": it["name"], "label": it.get("label", ""), "levels": level_info})
        if per_item:
            ins_name = next((i["name"] for i in instruments if i["id"] == cid), f"仪器{cid}")
            summary.append({"instrument_id": cid, "instrument_name": ins_name, "items": per_item})
    return {
        "category": "定量",
        "instruments": instruments,
        "compared": compared,
        "matrix": matrix,
        "summary": summary,
        "ref_id": ref_id,
        "item_eff_refs": item_eff_refs,
    }


# ---------------------------------------------------------------------------
# HTML 预览
# ---------------------------------------------------------------------------
def build_html(group: ComparisonGroup, plan: ComparisonPlan, data: dict):
    css = """
    <style>
      .rep { font-family: 'Microsoft YaHei','SimSun',serif; color:#222; max-width:960px; margin:0 auto; padding:16px; }
      .rep h1 { text-align:center; font-size:20px; margin:6px 0; }
      .rep .sub { text-align:center; font-size:12px; color:#555; margin-bottom:10px; }
      .rep h2 { font-size:14px; margin:14px 0 6px; border-left:4px solid #1a365d; padding-left:8px; }
      .rep p { font-size:13px; line-height:1.7; margin:4px 0; }
      .rep table { border-collapse:collapse; width:100%; font-size:12px; margin:6px 0; }
      .rep th, .rep td { border:1px solid #444; padding:4px 6px; text-align:center; }
      .rep th { background:#eef2f8; }
      .rep .item { text-align:left; }
      .rep .no { color:#c0392b; font-weight:700; }
      .rep .yes { color:#27ae60; font-weight:700; }
      .rep .mask { background:#f0f0f0; color:#aaa; }
      .rep .foot { margin-top:18px; font-size:13px; }
      .rep .summary-ok { color:#27ae60; }
    </style>"""
    title = group.form_title or ("定性室内比对结果记录及分析报告表" if group.category == "定性" else "定量室内比对结果记录分析表")
    form_code = group.form_code or ""
    year = plan.year or ""
    kind = "定性" if group.category == "定性" else "定量"
    html = [f'<div class="rep">{css}'
            f'<h1>{kind}室内比对报告（{group.name}）</h1>'
            f'<div class="sub" style="text-align:center;font-size:13px">民航总医院检验科　　生化免疫组　　比对日期：{plan.compared_at or ""}</div>'
            f'<div class="sub">表格编号 BG-SM-CZ-022（封面）　　{form_code}（数据页）　　民航总医院检验科生化免疫组　　生效日期：{year}.01.01</div>']

    # 第一页（022 封面）：比对方案 / 结果分析 / 处理方案
    ref = next((i for i in data["instruments"] if i["is_reference"]), None)
    compared_names = "、".join(i["name"] for i in data["instruments"] if not i["is_reference"])
    n_inst = len(data["instruments"])
    if data["category"] == "定性":
        items = [r["item"] for r in data["matrix"]]
        item_names = "、".join(items)
        n_item = len(items)
        plan_lines = [
            f"1.样本：{group.sample_desc or '5例临床样本'}。",
            f"2.仪器：{ref['name'] if ref else ''}、{compared_names}。",
            f"3.试验内容：将5例样本分别在{n_inst}台仪器上测定共有项目。",
            f"4.项目：{item_names}，共{n_item}项。",
            "5.评价标准：阴阳符合率≥80%代表比对可接受。",
        ]
    else:
        seen = set(); items = []
        for blk in data["matrix"]:
            for r in blk["rows"]:
                if r["item"] not in seen:
                    seen.add(r["item"]); items.append(r["item"])
        item_names = "、".join(items)
        n_item = len(items)
        has_abs = any(r["mode"] == "absolute" for blk in data["matrix"] for r in blk["rows"])
        plan_lines = [
            f"1.样本：{group.sample_desc or '5个不同浓度水平的室间质评样本'}。",
            f"2.仪器：{ref['name'] if ref else ''}、{compared_names}。",
            f"3.试验内容：将{group.levels or 5}份样本分别在{n_inst}台仪器上测定共有项目。",
            f"4.项目：{item_names}，共{n_item}项。",
            f"5.评价标准：允许偏倚参照行标/国家临检中心EQA评价准则，本实验室以{ref['name'] if ref else ''}结果为参照，"
            f"{'绝对' if has_abs else '相对'}偏倚绝对值应小于允许偏倚。是否可接受用Y/N表示。各项目成绩&gt;80%代表比对可接受。",
        ]
    html.append("<h2>比对方案</h2>")
    for line in plan_lines:
        html.append(f"<p>{line}</p>")
    html.append(f'<h2>结果分析</h2><p>{plan.summary or "各仪器上述所有项目均比对合格。"}</p>')
    html.append('<h2>处理方案（如不合格）</h2><p>{}</p>'.format(plan.handle_plan or "无"))

    # 后续页（数据页）
    html.append('<hr style="margin:18px 0"><h2>数据页：{title}（{form_code}）</h2>'.format(title=title, form_code=form_code))

    if data["category"] == "定性":
        html.append("<h2>比对结果</h2>")
        html.append('<table><thead><tr><th class="item">检验项目</th>')
        for ins in data["instruments"]:
            tag = "（参照）" if ins["is_reference"] else ""
            html.append(f'<th>{ins["name"]}{tag}</th><th>符合率</th>')
        html.append("</tr></thead><tbody>")
        for row in data["matrix"]:
            html.append(f'<tr><td class="item">{row["item"]}</td>')
            for ins in data["instruments"]:
                c = row["insts"].get(ins["id"], {})
                if c.get("masked"):
                    html.append('<td class="mask">/</td><td class="mask">/</td>')
                    continue
                cell = c.get("results") or []
                samples = " ".join(cell) if isinstance(cell, list) and cell else "-"
                agr = c.get("agreement")
                agr_s = f"{agr}%" if agr is not None else "-"
                html.append(f"<td>{samples}</td><td>{agr_s}</td>")
            html.append("</tr>")
        html.append("</tbody></table>")
        _qual_vals = [i["agreement"] for row in data["matrix"] for i in row["insts"].values()
                      if not i.get("masked")]
        _qual_ok = _qual_vals and all(a is not None and a >= 80 for a in _qual_vals)
        html.append(f'<p>总结：使用5例样本进行室内比对，结果阴阳符合率见上表，结果一致性{"可接受" if _qual_ok else "待评估"}。</p>')
    else:
        inst_name_map = {i["id"]: i["name"] for i in data["instruments"]}
        for blk in data["matrix"]:
            html.append(f"<h2>水平{blk['level']}</h2>")
            html.append('<table><thead><tr><th class="item">项目</th>'
                        '<th>参照值（靶机）</th><th>允许偏倚%</th>')
            for ins in data["instruments"]:
                if ins["is_reference"]:
                    continue
                html.append(f'<th>{ins["name"]}</th><th>偏倚%</th><th>是否允许</th>')
            html.append("</tr></thead><tbody>")
            for r in blk["rows"]:
                eff_name = inst_name_map.get(r.get("effective_ref_id"), "")
                ref_label = f'{r["ref"]}' if not eff_name else f'{r["ref"]}（{eff_name}）'
                html.append(f'<tr><td class="item">{r["item"]}</td><td>{ref_label}</td><td>{r["te"]}</td>')
                for ins in data["instruments"]:
                    if ins["is_reference"]:
                        continue
                    c = r["insts"].get(str(ins["id"]), {})
                    if c.get("masked"):
                        html.append('<td class="mask">/</td><td class="mask">/</td><td class="mask">/</td>')
                        continue
                    if c.get("is_ref"):
                        html.append(f'<td>{c.get("value","")}</td><td style="color:#409eff;font-weight:600">参照</td><td>—</td>')
                        continue
                    bias = c.get("bias"); acc = c.get("accepted")
                    bias = c.get("bias")
                    acc = c.get("accepted")
                    bias_s = f"{bias}%" if bias is not None else "-"
                    acc_cls = "yes" if acc is True else ("no" if acc is False else "")
                    acc_s = {True: "Y", False: "N", None: "-"}[acc]
                    html.append(f'<td>{c.get("value","")}</td><td>{bias_s}</td><td class="{acc_cls}">{acc_s}</td>')
                html.append("</tr>")
            html.append("</tbody></table>")

        html.append("<h2>汇总</h2>")
        for s in data["summary"]:
            html.append(f'<h3 style="font-size:13px;margin:8px 0 4px;color:#1a365d">▶ {s["instrument_name"]}（{len(s["items"])} 项）</h3>')
            html.append('<table><thead><tr><th class="item">项目</th>')
            for lv in range(1, (group.levels or 5) + 1):
                html.append(f"<th>水平{lv}</th>")
            html.append("</tr></thead><tbody>")
            for it in s["items"]:
                html.append(f'<tr><td class="item">{it["item"]}</td>')
                for info in it["levels"]:
                    ok = info["ok"]
                    cls = "yes" if ok else "no"
                    ref = str(info.get("ref") or "")
                    ref_line = f"{ref}<br>" if ref else ""
                    html.append(f'<td>{ref_line}<span class="{cls}">{"Y" if ok else "N"}</span></td>')
                html.append("</tr>")
            html.append("</tbody></table>")

    html.append(f'<div class="foot">操作者：{plan.operator or "　　　　"}　　审核者：{plan.reviewer or "　　　　"}　　日期：{plan.compared_at or "　　　　"}</div>')
    html.append("</div>")
    return "".join(html)


# ---------------------------------------------------------------------------
# DOCX 生成
# ---------------------------------------------------------------------------
def _set_cell_font(cell, size=10.5, bold=False, align="center"):
    for p in cell.paragraphs:
        p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.name = "SimSun"
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _fill(cell, text, size=10.5, bold=False, align="center", color=None):
    cell.text = str(text)
    _set_cell_font(cell, size, bold, align)
    if color:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = color


def _fill_ref_yn(cell, ref, ok, size=10.5, align="center"):
    """汇总单元格：上一行写靶机浓度，下一行写 Y/N（红绿色）。"""
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run1 = p1.add_run(str(ref) if ref not in (None, "") else "—")
    _run_font(run1, size, bold=False)
    p2 = cell.add_paragraph()
    p2.alignment = p1.alignment
    acc_s = "Y" if ok else "N"
    color = RGBColor(0x27, 0xae, 0x60) if ok else RGBColor(0xc0, 0x39, 0x2b)
    run2 = p2.add_run(acc_s)
    _run_font(run2, size, bold=True)
    run2.font.color.rgb = color


def _heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    return p


def _run_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    return run


def _add_field(run, field):
    """在 run 内插入一个 Word 域（PAGE / NUMPAGES）。"""
    f1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    instr = parse_xml(f'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> {field} </w:instrText>')
    fsep = parse_xml(r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    f2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    run._r.append(f1); run._r.append(instr); run._r.append(fsep); run._r.append(f2)


def _add_footer_code(section, code, year):
    """页脚：3 列 1 行表 → 左 表格编号/生效日期（堆 2 行），中 医院名，右 第X页/共Y页。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    # 清空默认段落
    for p in list(footer.paragraphs):
        p.text = ""
    # 1×3 表格做三栏
    tbl = footer.add_table(rows=1, cols=3, width=Cm(17.4))
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 去掉表格边框
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        borders.append(b)
    tblPr.append(borders)
    # 列宽：左 5cm / 中 7.4cm / 右 5cm
    widths = [Cm(5), Cm(7.4), Cm(5)]
    for i, w in enumerate(widths):
        for c in tbl.columns[i].cells:
            c.width = w
    left, center, right = tbl.rows[0].cells

    # 左：两行
    left.text = ""
    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(p1.add_run(f"表格编号：{code}"), 9)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(p2.add_run(f"生效日期：{year}.01.01"), 9)

    # 中：医院名
    center.text = ""
    p = center.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run("民航总医院检验科生化免疫组"), 9)

    # 右：第 X 页 共 Y 页
    right.text = ""
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run_font(p.add_run("第 "), 9)
    _add_field(p.add_run(""), "PAGE")
    _run_font(p.add_run(" 页 共 "), 9)
    _add_field(p.add_run(""), "NUMPAGES")
    _run_font(p.add_run(" 页"), 9)


def _cover_header(doc, plan):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(
        f"民航总医院检验科　　　　生化免疫组　　　　比对日期：{plan.compared_at or '　　　　　'}"), 10.5)


def _cover_title(doc, group):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kind = "定性" if group.category == "定性" else "定量"
    _run_font(p.add_run(f"{kind}室内比对报告（{group.name}）"), 16, bold=True)


def _cover_plan(doc, group, plan, data):
    ref = next((i for i in data["instruments"] if i["is_reference"]), None)
    compared = [i for i in data["instruments"] if not i["is_reference"]]
    compared_names = "、".join(i["name"] for i in compared)
    n_inst = len(data["instruments"])
    if group.category == "定性":
        items = [r["item"] for r in data["matrix"]]
        n_item = len(items)
        item_names = "、".join(items)
        lines = [
            f"1.样本：{group.sample_desc or '5例临床样本'}。",
            f"2.仪器：{ref['name'] if ref else ''}、{compared_names}。",
            f"3.试验内容：将5例样本分别在{n_inst}台仪器上测定共有项目。",
            f"4.项目：{item_names}，共{n_item}项。",
            "5.评价标准：阴阳符合率≥80%代表比对可接受。",
        ]
    else:
        seen = set(); items = []
        for blk in data["matrix"]:
            for r in blk["rows"]:
                if r["item"] not in seen:
                    seen.add(r["item"]); items.append(r["item"])
        n_item = len(items)
        item_names = "、".join(items)
        has_abs = any(r["mode"] == "absolute" for blk in data["matrix"] for r in blk["rows"])
        lines = [
            f"1.样本：{group.sample_desc or '5个不同浓度水平的室间质评样本'}。",
            f"2.仪器：{ref['name'] if ref else ''}、{compared_names}。",
            f"3.试验内容：将{group.levels or 5}份样本分别在{n_inst}台仪器上测定共有项目。",
            f"4.项目：{item_names}，共{n_item}项。",
            f"5.评价标准：允许偏倚参照行标/国家临检中心EQA评价准则，本实验室以{ref['name'] if ref else ''}结果为参照，"
            f"{'绝对' if has_abs else '相对'}偏倚绝对值应小于允许偏倚。是否可接受用Y/N表示。各项目成绩>80%代表比对可接受。",
        ]
    _heading(doc, "比对方案")
    for line in lines:
        p = doc.add_paragraph()
        _run_font(p.add_run(line), 11)


def _cover_analysis(doc, plan):
    _heading(doc, "结果分析")
    p = doc.add_paragraph()
    _run_font(p.add_run(plan.summary or "各仪器上述所有项目均比对合格。"), 11)


def _cover_handle(doc, plan):
    _heading(doc, "处理方案（如不合格）")
    p = doc.add_paragraph()
    _run_font(p.add_run(plan.handle_plan or "无"), 11)


def _data_page_header(doc, group, plan):
    title = group.form_title or ("定性室内比对结果记录及分析报告表" if group.category == "定性" else "定量室内比对结果记录分析表")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(title), 15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run(
        f"表格编号 {group.form_code or ''}　　民航总医院检验科生化免疫组　　生效日期：{plan.year or ''}.01.01"), 10.5)


def build_docx(db, group: ComparisonGroup, plan: ComparisonPlan, data: dict, out_path: str):
    doc = Document()
    year = plan.year or ""

    # 第一页：BG-SM-CZ-022 封面（含 比对方案/结果分析/处理方案，表格编号置页脚）
    sec0 = doc.sections[0]
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)
    _add_footer_code(sec0, "BG-SM-CZ-022", year)
    _cover_title(doc, group)
    _cover_header(doc, plan)
    _cover_plan(doc, group, plan, data)
    _cover_analysis(doc, plan)
    _cover_handle(doc, plan)

    # 后续页：024/025/026/027/071 数据页（仅详细结果表，独立分节+各自页脚编号）
    sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)
    _add_footer_code(sec1, group.form_code or "", year)
    _data_page_header(doc, group, plan)

    if data["category"] == "定性":
        insts = data["instruments"]
        cols = 1 + len(insts) * 2
        t = doc.add_table(rows=1, cols=cols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        _fill(hdr[0], "检验项目", bold=True)
        for i, ins in enumerate(insts):
            _fill(hdr[1 + i * 2], f'{ins["name"]}（参照）' if ins["is_reference"] else ins["name"], bold=True)
            _fill(hdr[2 + i * 2], "符合率", bold=True)
        all_ok = True
        for row in data["matrix"]:
            cells = t.add_row().cells
            _fill(cells[0], row["item"], align="left")
            for i, ins in enumerate(insts):
                c = row["insts"].get(ins["id"], {})
                if c.get("masked"):
                    _fill(cells[1 + i * 2], "/"); _fill(cells[2 + i * 2], "/")
                    continue
                res = c.get("results") or []
                samples = " ".join(res) if isinstance(res, list) and res else "-"
                agr = c.get("agreement")
                agr_s = f"{agr}%" if agr is not None else "-"
                if agr is not None and agr < 80:
                    all_ok = False
                _fill(cells[1 + i * 2], samples)
                _fill(cells[2 + i * 2], agr_s)
        p = doc.add_paragraph()
        _run_font(p.add_run(
            f"总结：使用5例样本进行室内比对，结果阴阳符合率见上表，结果一致性{'可接受' if all_ok else '需关注'}。"), 11)
    else:
        ref = next((i for i in data["instruments"] if i["is_reference"]), None)
        compared = [i for i in data["instruments"] if not i["is_reference"]]
        # 预建 {inst_id → name} 查找
        inst_name_map = {i["id"]: i["name"] for i in data["instruments"]}
        for blk in data["matrix"]:
            _heading(doc, f"水平{blk['level']}")
            ncol = 3 + len(compared) * 3
            t = doc.add_table(rows=1, cols=ncol)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            _fill(hdr[0], "项目", bold=True)
            _fill(hdr[1], "参照值（靶机）", bold=True)
            _fill(hdr[2], "允许偏倚%", bold=True)
            ci = 3
            for ins in compared:
                _fill(hdr[ci], ins["name"], bold=True); _fill(hdr[ci + 1], "偏倚%", bold=True); _fill(hdr[ci + 2], "是否允许", bold=True)
                ci += 3
            for r in blk["rows"]:
                cells = t.add_row().cells
                _fill(cells[0], r["item"], align="left")
                # 参照值：值 + 有效靶机名（如果与组参照不同则附加说明）
                eff_name = inst_name_map.get(r.get("effective_ref_id"), "")
                ref_label = f'{r["ref"]}' if not eff_name else f'{r["ref"]}（{eff_name}）'
                _fill(cells[1], ref_label)
                _fill(cells[2], r["te"])
                ci = 3
                for ins in compared:
                    c = r["insts"].get(str(ins["id"]), {})
                    if c.get("masked"):
                        _fill(cells[ci], "/"); _fill(cells[ci + 1], "/"); _fill(cells[ci + 2], "/")
                        ci += 3
                        continue
                    if c.get("is_ref"):
                        # 该项目以此仪器为靶机
                        _fill(cells[ci], c.get("value", ""));
                        _fill(cells[ci + 1], "参照", size=9, color=RGBColor(0x40, 0x9e, 0xff));
                        _fill(cells[ci + 2], "—", size=9)
                    bias = c.get("bias"); acc = c.get("accepted")
                    bias_s = f"{bias}%" if bias is not None else "-"
                    acc_s = {True: "Y", False: "N", None: "-"}[acc]
                    _fill(cells[ci], c.get("value", ""))
                    _fill(cells[ci + 1], bias_s)
                    col = RGBColor(0x27, 0xae, 0x60) if acc is True else (RGBColor(0xc0, 0x39, 0x2b) if acc is False else None)
                    _fill(cells[ci + 2], acc_s, color=col, bold=True)
                    ci += 3

        _heading(doc, "汇总")
        levels_n = group.levels or 5
        for s in data["summary"]:
            # 每台仪器：仪器名 + 项目×水平 表
            sub = doc.add_paragraph()
            sub.paragraph_format.space_before = Pt(6)
            sr = sub.add_run(f"▶ {s['instrument_name']}（{len(s['items'])} 项）")
            sr.font.size = Pt(11)
            sr.font.bold = True
            sr.font.name = "SimSun"
            sr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            t = doc.add_table(rows=1, cols=1 + levels_n)
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            _fill(hdr[0], "项目", bold=True)
            for lv in range(1, levels_n + 1):
                _fill(hdr[lv], f"水平{lv}", bold=True)
            for it in s["items"]:
                cells = t.add_row().cells
                _fill(cells[0], it["item"], align="left")
                for j, info in enumerate(it["levels"]):
                    _fill_ref_yn(cells[1 + j], info.get("ref"), info["ok"])

    # 数据页底部：操作者 / 审核者 / 日期
    foot = doc.add_paragraph()
    _run_font(foot.add_run(
        f"操作者：{plan.operator or '　　　　'}　　审核者：{plan.reviewer or '　　　　'}　　日期：{plan.compared_at or '　　　　'}"), 11)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# 定量比对结果 Excel 导入（解析 BG-SM-CZ-025 等"横排分水平"比对表）
# ---------------------------------------------------------------------------
_EXCEL_KEYWORDS = {"允许TE%", "允许TE", "偏倚%", "是否允许", "是否允许Y/N", "TE%", "偏倚"}


def _norm_token(s):
    """规范化字符串：去除非字母数字字符、转小写。安全处理 int/float。"""
    return _re.sub(r"[^a-z0-9]", "", str(s or "").lower())


def _is_instr_header(v):
    """Excel 表头行里，某单元格是否为仪器名列（而非 允许TE%/偏倚%/是否允许 等关键字，
    也非纯数字/百分比，也非'水平N'标记，也非 Y/N 判定列）。"""
    if v is None:
        return False
    s = str(v).strip()
    if s == "":
        return False
    if s in _EXCEL_KEYWORDS:
        return False
    if "水平" in s:
        return False
    if s in ("项目", "项目名称", "检验项目"):
        return False
    if s.upper() in ("Y", "N", "YN"):
        return False
    # 纯数字 / 百分比不是仪器列
    if s.replace(".", "").replace("-", "").replace("%", "").isdigit():
        return False
    return True


def _match_instrument(header, instruments):
    """header: Excel 列头（如 AU5821-B / 5811-急诊 / AU5822(标准) / HT7600）；
    instruments: [{id,name,model,is_reference}]。
    返回匹配到的系统仪器 id，或 None。"""
    h = str(header or "").strip()
    if not h:
        return None
    hn = _norm_token(h)
    # 精确（规范化后）
    for ins in instruments:
        if _norm_token(ins.get("name", "")) == hn or _norm_token(ins.get("model", "")) == hn:
            return ins["id"]
    # 包含"标准" → 参照仪器（模板中参照仪标记为 AU5822(标准) 等）
    if "标准" in h:
        for ins in instruments:
            if ins.get("is_reference"):
                return ins["id"]
    # 含"急诊" → 匹配含"急诊"或"急"的仪器（处理 DXI800-急诊 ↔ DXI800急 之类命名差异）
    if "急诊" in h or "急" in h:
        for ins in instruments:
            nm = ins.get("name", "") + ins.get("model", "")
            if "急诊" in nm or ("急" in nm and "急" in h):
                return ins["id"]
    # AU5822 → 实际上是 AU5821 系列（旧命名，已停用；新机为 AU5821 A/B）
    if "au5822" in hn:
        for ins in instruments:
            n = _norm_token(ins.get("name", "") + ins.get("model", ""))
            if "au5821" in n:
                return ins["id"]
    # HT7600 → 常见旧型号；优先匹配名字中含 ht 且非参照的仪器
    if "ht7600" in hn:
        for ins in instruments:
            if not ins.get("is_reference"):
                n = _norm_token(ins.get("name", "") + ins.get("model", ""))
                if "ht" in n or "7600" in n:
                    return ins["id"]
        # 实在找不到，返回第一个非参照仪
        for ins in instruments:
            if not ins.get("is_reference"):
                return ins["id"]
    # AU5821 + 后缀字母
    m = _re.search(r"au5821([ab])", hn)
    if m:
        L = m.group(1)
        for ins in instruments:
            n = _norm_token(ins.get("name", "") + ins.get("model", ""))
            if "au5821" in n and n.endswith(L):
                return ins["id"]
    # DXI800 + 序号（DXI800-3 → DXI800 3；保留原子串匹配作为兜底）
    m = _re.search(r"dxi800[ -]?(\d)", hn)
    if m:
        L = m.group(1)
        for ins in instruments:
            n = _norm_token(ins.get("name", "") + ins.get("model", ""))
            if n.endswith(L):
                return ins["id"]
    if "dxi800" in hn:
        for ins in instruments:
            if "dxi800" in _norm_token(ins.get("name", "") + ins.get("model", "")):
                return ins["id"]
    # TOP700 + 后缀字母
    if "top700" in hn:
        mm = _re.search(r"top700([abc])", hn)
        L = mm.group(1) if mm else None
        for ins in instruments:
            n = _norm_token(ins.get("name", "") + ins.get("model", ""))
            if "top700" in n and (L is None or n.endswith(L)):
                return ins["id"]
    # 通用子串兜底
    for ins in instruments:
        n = _norm_token(ins.get("name", "") + ins.get("model", ""))
        if hn and (hn in n or n in hn):
            return ins["id"]
    return None


# 定量比对 Excel 常用缩写 ↔ 系统项目代码 同义表。
# SOP 表单（BG-SM-CZ-025 等）用 BUN/CHOL/CRE，而系统分组用 UREA/TC/Cr 等代码，
# 二者需打通才能正确匹配。键/值均为规范化(去非字母数字+小写)后的名称。
QUANT_SYNONYMS = {
    "bun": "urea",           # Blood Urea Nitrogen = 尿素
    "chol": "tc",            # Cholesterol = 总胆固醇
    "cre": "cr",             # Creatinine = 肌酐
    # DXI800 模板常见缩写
    "fer": "ferr",           # Ferritin
    "folw": "叶酸",          # Folate（中文名）
    "hstni": "ctni",        # hsTnI → cTnI（高敏肌钙蛋白）
    "ckmb": "ck-mb",        # CK-MB
    "ifab": "ifa",           # Intrinsic Factor Antibody
    # 早孕模板
    "hcg": "β-hcg",         # HCG → β-HCG
    "孕酮": "prog",         # 孕酮 → Progesterone
    "雌二醇": "e2",         # 雌二醇 → Estradiol (E2)
}

# 反向同义：系统分组代码 → Excel 缩写（用于回退查 TE_LOOKUP 标准允许TE）
_GROUP_TO_EXCEL = {v: k for k, v in QUANT_SYNONYMS.items()}


def _resolve_te(it: dict, level: int = None):
    """允许TE 解析：优先用分组里配置好的 te；若为空/0，则回退到内置 TE_LOOKUP
    （含同义缩写映射，如分组 UREA → TE_LOOKUP 的 BUN=3）。保证缺配项目也能算出偏倚与判定。

    若提供 `level`，先查 `it["te_by_level"][str(level)]`，再回退到 `it["te"]`，再 TE_LOOKUP。
    """
    by_lv = it.get("te_by_level") or {}
    if level is not None and str(level) in by_lv and str(by_lv.get(str(level), "")).strip() not in ("", "0", "0.0"):
        return by_lv[str(level)]
    te = it.get("te")
    if str(te or "").strip() not in ("", "0", "0.0"):
        return te
    name = _norm_token(it.get("name", ""))
    up = name.upper()
    if up in TE_LOOKUP:
        return TE_LOOKUP[up][0]
    ex = _GROUP_TO_EXCEL.get(name)
    if ex and ex.upper() in TE_LOOKUP:
        return TE_LOOKUP[ex.upper()][0]
    return te or "0"


def _resolve_mode(it: dict, level: int = None, default: str = "relative"):
    """偏倚方式解析：先 `mode_by_level[str(level)]` > `mode` > default。"""
    by_lv = it.get("mode_by_level") or {}
    if level is not None and str(level) in by_lv and by_lv[str(level)]:
        return by_lv[str(level)]
    m = it.get("mode")
    return m if m else default


def _match_item(name, items):
    """name: Excel 项目名；items: [{name,label,...}]（系统分组项目）。按规范化代码匹配，
    并支持 QUANT_SYNONYMS 同义缩写。"""
    h = str(name or "").strip()
    if not h:
        return None
    hn = _norm_token(h)
    # 1) 精确（规范化后）
    for it in items:
        if _norm_token(it.get("name", "")) == hn:
            return it["name"]
    # 2) 同义缩写映射（Excel 缩写 → 系统代码）
    syn = QUANT_SYNONYMS.get(hn)
    if syn:
        for it in items:
            # 先试规范化（英文字母缩写）
            if _norm_token(it.get("name", "")) == syn:
                return it["name"]
            # 中文字段名直接比较（如 叶酸 → syn="叶酸"）
            if it.get("name", "") == syn:
                return it["name"]
    # 3) 子串兜底
    for it in items:
        if hn and hn in _norm_token(it.get("name", "")):
            return it["name"]
    return None


def import_quant_from_excel(db, group: ComparisonGroup, plan: ComparisonPlan, file_bytes: bytes):
    """解析定量比对结果 Excel，返回 {quant:[...], levels, matched_items, unmatched_items, instruments_matched}。

    适用布局：表头行含"允许TE%"，每个水平块内仪器列依次是 [参照仪器, 允许TE%, 仪器B, 偏倚%, 是否允许, 仪器C, ...]；
    项目名在 A 列；一个 sheet 可含多个水平块（水平1/2/3 横排，4/5 在下方另一表头行）。
    """
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(file_bytes), data_only=True)
    instruments = _instruments_of(db, group)  # [{id,name,model,is_reference}]
    items = _load_json(group.items, [])
    K = max(len(instruments), 1)

    quant = []
    matched_items = set()
    unmatched_items = []
    instr_matched = set()

    for ws in wb.worksheets:
        maxr, maxc = ws.max_row, ws.max_column
        # 表头行唯一可靠标志：含"允许TE%"或"偏倚%"列（数据行的 Y/N 不会误判）
        header_rows = [
            r for r in range(1, maxr + 1)
            if any(str(ws.cell(r, c).value or "").strip() in ("允许TE%", "偏倚%", "允许TE")
                   for c in range(1, maxc + 1))
        ]
        level_counter = 0
        for H in sorted(set(header_rows)):
            instr_cols = [(c - 1, ws.cell(H, c).value)
                         for c in range(1, maxc + 1) if _is_instr_header(ws.cell(H, c).value)]
            if not instr_cols:
                continue
            chunks = [instr_cols[i:i + K] for i in range(0, len(instr_cols), K)]
            for chunk in chunks:
                level_counter += 1
                level = level_counter
                for r in range(H + 1, maxr + 1):
                    nm = ws.cell(r, 1).value
                    if nm is None or str(nm).strip() == "":
                        break
                    if "水平" in str(nm):
                        break
                    if r in set(header_rows):
                        break
                    mitem = _match_item(nm, items)
                    if not mitem:
                        s = str(nm).strip()
                        if s not in unmatched_items:
                            unmatched_items.append(s)
                        continue
                    matched_items.add(mitem)
                    vals = {}
                    for (col, hdr) in chunk:
                        iid = _match_instrument(hdr, instruments)
                        if iid is None:
                            continue
                        # instr_cols 中的 col 为 0 基，+1 即回到 1 基的表头同列
                        # （BG-SM-CZ-025 布局：仪器表头与对应测量值同列）
                        v = ws.cell(r, col + 1).value
                        if v is None or str(v).strip() == "":
                            continue
                        vals[iid] = v
                        instr_matched.add(iid)
                    if not vals:
                        continue
                    ref_id = group.reference_instrument_id
                    ref_val = vals.get(ref_id)
                    if ref_val is None and vals:
                        ref_id = next(iter(vals.keys()))
                        ref_val = vals[ref_id]
                    compared = {iid: v for iid, v in vals.items() if iid != ref_id}
                    quant.append({
                        "item": mitem,
                        "level": level,
                        "reference_value": "" if ref_val is None else str(ref_val),
                        "values": {str(k): ("" if v is None else str(v)) for k, v in compared.items()},
                    })

    return {
        "quant": quant,
        "levels": max((q["level"] for q in quant), default=0),
        "matched_items": sorted(matched_items),
        "unmatched_items": unmatched_items,
        "instruments_matched": sorted(instr_matched),
    }


# ---------------------------------------------------------------------------
# 默认分组种子（参照 SOP 表单 BG-SM-CZ-021/024~027/071）
# 允许偏倚(TE) 默认取自 WS/T 403—2024《临床化学检验常用项目分析质量标准》附录 A
#  资料来源：C:\Users\81526\Desktop\WST403-2024临床化学检验常用项目分析质量标准.pdf
#  该标准为推荐性国标（2024-05-09 发布 2024-11-01 实施，代替 WS/T 403—2012）
#  各项的"允许偏倚"列即为系统 te 字段的权威值（相对%）。
#  标准里允许总误差(TEa)对低浓度水平有"绝对值/相对%"分段公式，
#  对应"低浓度用绝对偏倚、高浓度用相对偏倚"——本系统已在 GroupItem
#  的 te_by_level / mode_by_level 字段里支持该场景。
# ---------------------------------------------------------------------------
def _items_quant(pairs):
    return [{"name": n, "te": str(te), "mode": m} for n, te, m in pairs]


# WS/T 403-2024 附录 A 允许偏倚（系统项目代码 → 偏倚%）。
# mode='absolute' 表示标准对该项目在所有水平都用绝对偏倚（如 pH）；
# mode='relative' 表示相对%（绝大多数项目）；低浓度水平可通过 te_by_level/mode_by_level 切换绝对偏倚。
WST403_2024 = {
    # ─ 常规生化 ─
    "K":  2.0,  "NA": 1.5,  "CL": 1.5,  "CA": 2.0,  "P":  3.0,  "GLU": 2.0,
    "UREA": 3.0,  "UA": 4.5,  "CR": 5.5,  "TP": 2.0,  "ALB": 2.0,  "TC": 4.0,
    "TG":  5.0,  "HDL": 8.0,  "LDL": 8.0,  "APOA1": 10.0, "APOB": 10.0, "LPA": 10.0,
    "TBIL": 5.0,  "DBIL": 6.7,  "ALT": 5.0,  "AST": 5.0,  "ALP": 10.0,  "AMY": 7.5,
    "CK": 5.5,  "LDH": 4.0,  "GGT": 5.5,  "HBDH": 10.0,
    "CHE": 8.0,  "FE": 4.5,  "MG": 5.5,  "CYSC": 8.0,
    # ─ 免疫 ─
    "CKMB": 10.0,  "MYO": 10.0,  "HCY": 10.0,
    "IGG": 8.0,  "IGA": 8.0,  "IGM": 10.0,
    "CRP": 10.0,  "RF": 10.0,  "ASO": 10.0,  "PA": 10.0,
    "FT3": 8.0,  "TT3": 8.0,  "FT4": 8.0,  "TT4": 8.0,  "TSH": 8.0,
    "CORTISOL": 8.0,  "E2": 10.0,  "FSH": 8.0,  "LH": 8.0,
    "PROG": 8.0,  "PRL": 8.0,  "TESTO": 8.0,  "INSULIN": 12.0,
    "FA": 12.0,  "B12": 10.0,  "PTH": 10.0,
    "FERR": 10.0,  "B2MG": 10.0,
    # ─ 血气（pH 绝对，其余相对%） ─
    "PH":   (0.015, "absolute"),  "PCO2": (4.0, "relative"),  "PO2": (5.0, "relative"),
}


def _wst(item_code: str, default_te: float = 10.0, default_mode: str = "relative"):
    """从 WS/T 403-2024 取允许偏倚；标准未列出的项目用 default_* 回退。"""
    v = WST403_2024.get(item_code.upper())
    if v is None:
        return default_te, default_mode
    if isinstance(v, tuple):
        return v
    return v, "relative"


BIOTH_ITEM = _items_quant([
    # 以下按 WS/T 403-2024 附录 A 校准（先前错误值已纠正）：
    ("ALB", *_wst("ALB")),       # 白蛋白
    ("ALP", *_wst("ALP")),       # 碱性磷酸酶
    ("ALT", *_wst("ALT")),       # 丙氨酸氨基转移酶
    ("AMY", *_wst("AMY")),       # 淀粉酶
    ("APA", *_wst("APOA1")),     # 载脂蛋白 A1（系统别名 APA → APOA1）
    ("APB", *_wst("APOB")),      # 载脂蛋白 B
    ("ASO", *_wst("ASO")),       # 抗链球菌溶血素 O
    ("AST", *_wst("AST")),       # 天门冬氨酸氨基转移酶
    ("BUN", *_wst("UREA")),      # 尿素（Excel 缩写 BUN → 系统 UREA）
    ("C3",  8, "relative"),      # 补体 C3（标准未单列，沿用行业通用 8%）
    ("C4",  10, "relative"),     # 补体 C4
    ("CA", *_wst("CA")),
    ("CHE", *_wst("CHE")),       # 胆碱酯酶
    ("CHOL", *_wst("TC")),       # 总胆固醇
    ("CK", *_wst("CK")),
    ("CL", *_wst("CL")),
    ("CO2", 3, "relative"),      # 二氧化碳（标准未单列，沿用 3%）
    ("CRE", *_wst("CR")),        # 肌酐
    ("CRP", *_wst("CRP")),
    ("CYSC", *_wst("CYSC")),     # 胱抑素 C
    ("DBIL", *_wst("DBIL")),
    ("FE", *_wst("FE")),
    ("GA", 10, "relative"),      # 糖化白蛋白（标准未单列）
    ("GGT", *_wst("GGT")),
    ("GLU", *_wst("GLU")),
    ("HCY", *_wst("HCY")),
    ("HDL", *_wst("HDL")),
    ("Hp", 12, "relative"),      # 触珠蛋白（标准未单列）
    ("Ig-A", *_wst("IGA")),      # ← 之前错为 10，标准 8%
    ("Ig-G", *_wst("IGG")),
    ("Ig-M", *_wst("IGM")),      # ← 之前错为 8，标准 10%
    ("K", *_wst("K")),
    ("LAC", 12, "relative"),     # 乳酸（标准未单列）
    ("LDH", *_wst("LDH")),
    ("LDL", *_wst("LDL")),
    ("LPa", *_wst("LPA")),
    ("LPS", 10, "relative"),     # 脂肪酶（标准未单列）
    ("Mg", *_wst("MG")),
    ("NA", *_wst("NA")),
    ("P", *_wst("P")),
    ("PA", *_wst("PA")),
    ("RF", *_wst("RF")),
    ("SAA", 15, "relative"),     # 血清淀粉样蛋白 A（标准未单列）
    ("sd-LDL", 10, "relative"),  # 小而密 LDL（标准未单列）
    ("TBA", 10, "relative"),     # 总胆汁酸（标准未单列，10%）
    ("TBIL", *_wst("TBIL")),
    ("TG", *_wst("TG")),
    ("TP", *_wst("TP")),
    ("UA", *_wst("UA")),
    ("UIBC", 10, "relative"),    # 不饱和铁结合力（标准未单列）
    ("Zn", 10, "relative"),      # 锌（标准未单列）
    ("β2mg", *_wst("B2MG")),     # β2-微球蛋白
    ("β-HBDH", *_wst("HBDH")),   # ← 之前错为 15，标准 10%
])

DXI_ITEM = _items_quant([
    # 之前全部填 0.1 是占位符，现按 WS/T 403-2024 校准：
    ("FERR",  *_wst("FERR")),    # 铁蛋白
    ("叶酸",   *_wst("FA")),      # 叶酸
    ("B12",   *_wst("B12")),     # 维生素 B12
    ("sTfR",  10, "relative"),   # 可溶性转铁蛋白受体（标准未单列）
    ("IFA",   10, "relative"),   # 铁蛋白（已在 FERR 中；此处为去铁铁蛋白？沿用 10%）
    ("PCT",   10, "relative"),   # 降钙素原（标准未单列）
    ("IL-6",  12, "relative"),   # 白介素 6（标准未单列）
    ("cTnI",  10, "relative"),   # 心肌肌钙蛋白 I（标准未单列）
    ("MYO",   *_wst("MYO")),
    ("CK-MB", *_wst("CKMB")),    # 肌酸激酶-MB（μg/L，质量法）
    ("BNP",   10, "relative"),   # B 型利钠肽（标准未单列）
])

COAG_ITEM = _items_quant([
    ("D-D", 25, "relative"), ("APTT", 8, "relative"), ("AT-III", 8, "relative"),
    ("TT", 10, "relative"), ("PT", 8, "relative"), ("FDP", 7, "relative"), ("FIB", 10, "relative"),
])

PREG_ITEM = _items_quant([
    # 之前全部填 25 偏高，按 WS/T 403-2024 校准（β-HCG 标准未单列，沿用 10%）：
    ("HCG",   10, "relative"),   # β-HCG 标准未列
    ("孕酮",  *_wst("PROG")),    # 孕酮
    ("雌二醇", *_wst("E2")),     # 雌二醇
])

BLOODGAS_ITEM = [
    # 按 WS/T 403-2024 校准：
    {"name": "PH",   "te": "0.015", "mode": "absolute"},  # pH 标准为绝对偏倚 0.015
    {"name": "PCO2", "te": "4",     "mode": "relative"},  # CO2 分压 4%
    {"name": "pO2",  "te": "5",     "mode": "relative"},  # O2 分压 5%（原误为 0.05/相对，按标准改为 5%）
    {"name": "Na+",  "te": "1.5",   "mode": "relative"},  # 与血 Na 共享 1.5%
    {"name": "K+",   "te": "2",     "mode": "relative"},  # 与血 K 共享 2%
    {"name": "Ca2+", "te": "2",     "mode": "relative"},  # 与血 Ca 共享 2%
    {"name": "Cl-",  "te": "1.5",   "mode": "relative"},  # 与血 Cl 共享 1.5%
]

QUAL_ITEM = [
    {"name": "乙肝表面抗原", "te": "0", "mode": "relative"},
    {"name": "乙肝表面抗体", "te": "0", "mode": "relative"},
    {"name": "乙肝e抗原", "te": "0", "mode": "relative"},
    {"name": "乙肝e抗体", "te": "0", "mode": "relative"},
    {"name": "乙肝核心抗体", "te": "0", "mode": "relative"},
    {"name": "丙肝抗体", "te": "0", "mode": "relative"},
    {"name": "梅毒抗体", "te": "0", "mode": "relative"},
    {"name": "HIV抗体", "te": "0", "mode": "relative"},
]


# ---------------------------------------------------------------------------
# 从仪器档案解析共有项目（项目↔仪器关联）+ 每项适用仪器（用于遮蔽）
# ---------------------------------------------------------------------------
def _build_te_lookup():
    """合并各专业种子的允许TE，按项目代码(大写)索引，供自动填充默认TE。"""
    lut = {}
    for tbl in (BIOTH_ITEM, DXI_ITEM, COAG_ITEM, PREG_ITEM, BLOODGAS_ITEM):
        for it in tbl:
            lut[it["name"].upper()] = (it["te"], it["mode"])
    return lut


TE_LOOKUP = _build_te_lookup()


def _split_tokens(s: str):
    """把 instrument_group 文本拆成子型号 token，如 'AU58-1 / AU58-2/AU5800'。"""
    if not s:
        return []
    out = []
    for part in str(s).replace("／", "/").split("/"):
        t = part.strip()
        if t:
            out.append(t)
    return out


def resolve_common_items(db, instrument_ids, category="定量", min_count=2):
    """依据仪器档案里"项目↔仪器"关联，计算给定仪器组的共有项目。

    关联链：
    - test_items.instrument_group 里的子型号 token（如 AU58-1/AU5800）→ 小型号
      instrument_families(name) → members → 具体仪器 id；据此得到"每个项目实际装机的仪器"。
    - 若 instrument_group 为空，则回退用 test_items.instrument（总型号）→ family → members。

    返回：{items:[{name(代码), label(中文), te, mode, instrument_ids:[适用且在本组内]}],
           instruments:[{id,name,model}]}，仅保留在本组内适用仪器数 ≥ min_count 的项目。
    """
    from ..models.test_item import TestItem
    from ..models.instrument_family import InstrumentFamily, InstrumentFamilyMember

    cand = [int(i) for i in (instrument_ids or [])]
    cand_set = set(cand)
    if not cand_set:
        return {"items": [], "instruments": []}

    # family.name -> set(instrument_id)
    fam_name_to_ids = {}
    fam_id_to_name = {f.id: (f.name or "").strip() for f in db.query(InstrumentFamily).all()}
    for m in db.query(InstrumentFamilyMember).all():
        nm = fam_id_to_name.get(m.family_id, "")
        if nm:
            fam_name_to_ids.setdefault(nm, set()).add(m.instrument_id)

    def match_token_ids(token):
        """token 精确匹配小型号 family 名，返回其成员 id 集合。"""
        token = token.strip()
        if token in fam_name_to_ids:
            return set(fam_name_to_ids[token])
        # 宽松：忽略大小写/空格
        tl = token.lower().replace(" ", "")
        for nm, ids in fam_name_to_ids.items():
            if nm.lower().replace(" ", "") == tl:
                return set(ids)
        return set()

    items_out = []
    seen_names = set()
    # 1) 先检索现有已生成的报告（同 group 内、或跨 group），从已知 group items 里取真实用过的 te/mode 作为默认
    from ..models.comparison import ComparisonGroup
    used_default = {}  # code_upper -> {te, mode}
    for g in db.query(ComparisonGroup).all():
        for gi in _load_json(g.items, []):
            c = (gi.get("name") or "").strip().upper()
            if not c: continue
            cur = used_default.get(c)
            cand_te = gi.get("te"); cand_mode = gi.get("mode")
            if cur is None:
                used_default[c] = {"te": cand_te, "mode": cand_mode}
            else:
                # 保留最具体的：te/mode 非空时优先
                if (not cur.get("te")) and cand_te: cur["te"] = cand_te
                if (not cur.get("mode")) and cand_mode: cur["mode"] = cand_mode
    for ti in db.query(TestItem).all():
        # 该项目装机仪器 id 集合
        run_ids = set()
        for tok in _split_tokens(ti.instrument_group):
            run_ids |= match_token_ids(tok)
        if not run_ids:
            # 回退：总型号 family
            run_ids |= fam_name_to_ids.get((ti.instrument or "").strip(), set())
        app = sorted(run_ids & cand_set)
        if len(app) < min_count:
            continue
        code = (ti.code or "").strip() or (ti.name or "").strip()
        label = (ti.name or "").strip() or code
        key = code.upper()
        if key in seen_names:
            continue
        seen_names.add(key)
        # 默认 TE/mode 优先级：现有报告里已用过的 > TE_LOOKUP
        prior = used_default.get(key) or {}
        te_lk, mode_lk = TE_LOOKUP.get(key, ("", "relative"))
        te = prior.get("te") or te_lk or "0"
        mode = prior.get("mode") or mode_lk or "relative"
        items_out.append({
            "name": code, "label": label, "te": str(te), "mode": mode,
            "instrument_ids": app,
        })

    insts = []
    for iid in cand:
        inst = db.get(Instrument, iid)
        if inst:
            insts.append({"id": inst.id, "name": disp_name(inst), "model": inst.model or ""})
    return {"items": items_out, "instruments": insts}


def _find_instrument(db, *keywords, used=None):
    used = used or set()
    for inst in db.query(Instrument).all():
        nm = f"{inst.name or ''} {inst.model or ''}"
        for kw in keywords:
            if kw and kw.lower() in nm.lower() and inst.id not in used:
                return inst.id
    return 0


def ensure_comparison_defaults(db):
    """种子比对分组模板（幂等）。仪器按型号匹配在用机器且互不重复；
    定量组的项目清单与"每项适用仪器"由仪器档案自动解析（含遮蔽），解析不到时回退固定清单。"""
    seeds = [
        ("生化分析仪", "定量", "BG-SM-CZ-025", "定量室内比对结果记录分析表（生化分析仪）",
         ("AU5821 A", "AU5821 B", "AU5800急诊"), "AU5821 A", 5, BIOTH_ITEM,
         "5个不同浓度水平的室间质评样本"),
        ("DXI800分析仪", "定量", "BG-SM-CZ-024", "定量室内比对结果记录分析表（DXI800分析仪）",
         ("DXI800 1", "DXI800 2", "DXI800 3", "DXI800 4"), "DXI800 1", 5, DXI_ITEM,
         "5个不同浓度水平的新鲜血清样本"),
        ("凝血分析仪", "定量", "BG-SM-CZ-026", "定量室内比对结果记录分析表（凝血分析仪）",
         ("TOP700A", "TOP700B", "TOP700C"), "TOP700A", 5, COAG_ITEM,
         "5个不同浓度水平的新鲜血浆样本（至少包含2个凝血结果异常的样本）"),
        ("早孕系列", "定量", "BG-SM-CZ-027", "定量室内比对结果记录分析表（早孕系列）",
         ("DXI800 唐", "DXI800 急"), "DXI800 唐", 5, PREG_ITEM,
         "5个不同浓度水平的新鲜血清样本"),
        ("血气分析仪", "定量", "BG-SM-CZ-071", "定量室内比对结果记录分析表（血气分析仪）",
         ("RapidPoint1", "RapidPoint2"), "RapidPoint1", 2, BLOODGAS_ITEM,
         "2个浓度水平的血气专用质控物"),
        ("定性比对", "定性", "BG-SM-CZ-021", "定性室内比对结果记录及分析报告表",
         ("e601", "e602", "e411"), "e601", 0, QUAL_ITEM, "5例临床样本"),
    ]
    # 在用仪器 id 集合（用于判断旧种子是否失效）
    inuse = {i.id for i in db.query(Instrument).all()
             if (getattr(i, "status", "") or "") == "在用"}

    def _stale(ids_list):
        """旧的问题种子特征：有重复 id、为空、或含非在用仪器。"""
        if not ids_list:
            return True
        if len(ids_list) != len(set(ids_list)):
            return True
        if any(i not in inuse for i in ids_list):
            return True
        return False

    for name, cat, code, ftitle, inst_kw, ref_kw, levels, fallback_items, sample in seeds:
        existing = db.query(ComparisonGroup).filter_by(name=name).first()
        # 已存在且数据正常（非旧问题种子）→ 保留用户数据，不覆盖
        if existing and not _stale(_load_json(existing.instrument_ids, [])):
            continue

        # 计算正确的仪器组（按型号匹配在用机器且互不重复）
        used = set()
        ids = []
        for kw in inst_kw:
            iid = _find_instrument(db, kw, used=used)
            if iid:
                ids.append(iid)
                used.add(iid)
        ref_id = 0
        if ref_kw:
            ref_id = _find_instrument(db, ref_kw, used=used) or (ids[0] if ids else 0)
            if ref_id and ref_id not in ids:
                ids.insert(0, ref_id)
        # 定量组：优先用档案解析项目（带每项适用仪器/遮蔽）
        items = fallback_items
        if cat == "定量" and ids:
            resolved = resolve_common_items(db, ids, category=cat, min_count=2).get("items", [])
            if resolved:
                items = resolved

        if existing:
            # 修复旧问题种子（重复/失效 id）——仅刷新模板字段，保留 id 与关联计划
            existing.instrument_ids = json.dumps(ids)
            existing.reference_instrument_id = ref_id
            existing.items = json.dumps(items)
            existing.form_code = code
            existing.form_title = ftitle
            existing.sample_desc = sample
            existing.levels = levels
        else:
            db.add(ComparisonGroup(
                name=name, category=cat, form_code=code, form_title=ftitle,
                instrument_ids=json.dumps(ids), reference_instrument_id=ref_id,
                levels=levels, items=json.dumps(items), sample_desc=sample,
                created_by="system",
            ))
    db.commit()
