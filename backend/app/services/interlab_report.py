"""室间比对：结果计算 + 报告生成（docx）+ HTML 预览 + 候选项目解析。

模板参照 BG-SM-CZ-019（定量）/ BG-SM-CZ-018（定性）。
每个项目测 5 个水平的样本，每个水平录我室值 X + 比较系统1/2 的 Y1/Y2/均值Y。
"""

import json
import os
import calendar
from datetime import datetime, date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from ..models.instrument import Instrument
from ..models.test_item import TestItem

REPORT_DIR = None  # 由 main 注入（DATA_DIR/interlab_reports）


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def _parse_float(x):
    try:
        return float(str(x).strip())
    except Exception:
        return None


def _parse_date_str(s):
    """解析日期字符串，支持 ISO 日期及常见分隔符。"""
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(str(s).strip(), fmt).date()
        except Exception:
            pass
    return None


def _add_months(d: date, months: int) -> date:
    """给 date 增加指定月数，月末自动处理（1.31 + 1月 → 2.28/29）。"""
    total = d.month - 1 + months
    new_year = d.year + total // 12
    new_month = total % 12 + 1
    last_day = calendar.monthrange(new_year, new_month)[1]
    new_day = min(d.day, last_day)
    return date(new_year, new_month, new_day)


def _half_year_end(year: int, half: int) -> date:
    return date(year, 6, 30) if half == 1 else date(year, 12, 31)


def _fmt_before(d: date) -> str:
    return f"{d.year}年{d.month}月{d.day}日之前"


def _compute_bias(our, ref, te, mode):
    """返回 (偏倚值或None, 是否合格或None)。mode: relative(相对%) / absolute(绝对)"""
    of = _parse_float(our)
    rf = _parse_float(ref)
    tf = _parse_float(te)
    if of is None or rf is None or tf is None:
        return None, None
    if mode == "absolute":
        bias = of - rf
    else:
        if rf == 0:
            return None, None
        bias = (of - rf) / rf * 100.0
    accepted = abs(bias) <= tf + 1e-9
    return round(bias, 2), accepted


def _norm_pn(v):
    """阴阳判定归一化。
    支持完整词（"阳性"/"阴"/"POS" 等）和带括号/数值后缀的复合形式
    （"阴性(1.05)" / "阳性(6.78)" / "阴(+)" 等），从字符串前缀提取 阴/阳。"""
    s = (str(v or "").strip()).upper()
    if not s:
        return None
    pos_full = {"阳", "阳性", "POS", "POSITIVE", "P", "+", "1"}
    neg_full = {"阴", "阴性", "NEG", "NEGATIVE", "N", "-", "0"}
    if s in pos_full:
        return "positive"
    if s in neg_full:
        return "negative"
    # 前缀匹配：长前缀优先（"阳性" 在 "阳" 之前；"阴性" 在 "阴" 之前）
    prefixes = [
        ("阳性", "positive"), ("POSITIVE", "positive"), ("POS", "positive"),
        ("阳", "positive"), ("+", "positive"), ("P", "positive"),
        ("阴性", "negative"), ("NEGATIVE", "negative"), ("NEG", "negative"),
        ("阴", "negative"), ("-", "negative"), ("N", "negative"),
    ]
    for tok, kind in prefixes:
        if s.startswith(tok):
            return kind
    return None


def _split_tokens(s: str):
    if not s:
        return []
    out = []
    for part in str(s).replace("／", "/").split("/"):
        t = part.strip()
        if t:
            out.append(t)
    return out


def _family_name_to_ids(db):
    from ..models.instrument_family import InstrumentFamily, InstrumentFamilyMember
    fam_id_to_name = {f.id: (f.name or "").strip() for f in db.query(InstrumentFamily).all()}
    lut = {}
    for m in db.query(InstrumentFamilyMember).all():
        nm = fam_id_to_name.get(m.family_id, "")
        if nm:
            lut.setdefault(nm, set()).add(m.instrument_id)
    return lut


# ---------------------------------------------------------------------------
# 候选 / 必做项目
# ---------------------------------------------------------------------------
def candidate_projects(db, instrument_id: int):
    """返回该仪器可参与室间比对的项目（has_eqa=0 且 has_interlab=1）。"""
    inst = db.get(Instrument, instrument_id) if instrument_id else None
    if not inst:
        return []
    from ..models.instrument_family import InstrumentFamily, InstrumentFamilyMember
    fam_lut = _family_name_to_ids(db)
    fam_ids_names = {f.id: (f.name or "").strip() for f in db.query(InstrumentFamily).all()}
    inst_families: dict = {}
    for m in db.query(InstrumentFamilyMember).all():
        inst_families.setdefault(m.instrument_id, set()).add(m.family_id)
    fam_names = {fam_ids_names.get(fid, "") for fid in inst_families.get(instrument_id, set())}
    fam_names.discard("")
    extra_tokens = set()
    for t in _split_tokens(inst.name or ""):
        extra_tokens.add(t)
    if inst.model:
        extra_tokens.add(inst.model.strip())

    out = []
    seen = set()
    for ti in db.query(TestItem).all():
        he = getattr(ti, "has_eqa", None)
        if he is not None and int(he) == 1:
            continue
        hi = getattr(ti, "has_interlab", None)
        if hi is not None and int(hi) != 1:
            continue
        run_ids = set()
        for tok in _split_tokens(ti.instrument_group):
            run_ids |= fam_lut.get(tok, set())
        if not run_ids:
            run_ids |= fam_lut.get((ti.instrument or "").strip(), set())
        expanded = run_ids
        hit = (instrument_id in expanded) \
            or any(fn in _split_tokens(ti.instrument_group) for fn in fam_names) \
            or any(tok in (ti.instrument or "") for tok in extra_tokens if tok)
        if not hit:
            continue
        key = (ti.code or ti.name or "").strip().upper()
        if key in seen:
            continue
        seen.add(key)
        out.append({
            "id": ti.id,
            "code": (ti.code or "").strip(),
            "name": (ti.name or "").strip(),
            "unit": (ti.unit or "").strip(),
            "instrument": (ti.instrument or "").strip(),
            "mandatory": True,
        })
    return out


def mandatory_projects(db):
    """指导用：返回所有「无室间质评、需做室间比对」的必做项目及所属仪器。

    过滤逻辑：
    - has_eqa 改用 eqa_associations 的计算结果（_compute_associations），
      避免依赖 test_items.has_eqa 这个陈旧缓存字段。
    - has_interlab 仍用 test_items 字段（无变动）。

    进度：
    - last_plan：该项目最近一次被哪个 plan 覆盖（按 year/half/instrument 维度）。
    - last_status：最近一次 plan 的状态（done=已完成，其它=进行中/草稿）。
    - next_due：根据「每半年一次」规则推断下一次应做时间。
    """
    from ..api.v1.eqa_associations import _compute_associations
    fam_lut = _family_name_to_ids(db)
    inst_lut = {i.id: i for i in db.query(Instrument).all()}

    # 计算 EQA 关联（用实际 plan 数据）
    assoc = _compute_associations(db)
    has_eqa_map = {a.id: a.has_eqa for a in assoc}

    # 找每个 test_item 最近一次 plan（按名称匹配 InterlabItem.item）
    from ..models.interlab import InterlabPlan, InterlabItem
    # 一次性拿全部 items + plans
    plan_map = {p.id: p for p in db.query(InterlabPlan).all()}
    # item_name -> 最新 plan（按 year desc, half desc, plan.id desc）
    latest_plan_by_name: dict[str, InterlabPlan] = {}
    for it in db.query(InterlabItem).all():
        p = plan_map.get(it.plan_id)
        if not p:
            continue
        prev = latest_plan_by_name.get(it.item)
        if (not prev or
            (p.year, p.half, p.id) > (prev.year, prev.half, prev.id)):
            latest_plan_by_name[it.item] = p

    out = []
    for ti in db.query(TestItem).all():
        # 用实时计算的 has_eqa（不再用 ti.has_eqa 缓存）
        if has_eqa_map.get(ti.id):
            continue
        hi = getattr(ti, "has_interlab", None)
        if hi is not None and int(hi) != 1:
            continue

        run_ids = set()
        for tok in _split_tokens(ti.instrument_group):
            run_ids |= fam_lut.get(tok, set())
        if not run_ids:
            run_ids |= fam_lut.get((ti.instrument or "").strip(), set())
        insts = []
        for iid in sorted(run_ids):
            inst = inst_lut.get(iid)
            if not inst:
                continue
            if inst.status and "停用" in (inst.status or ""):
                continue
            nm = (inst.model or inst.name or "").strip() or f"仪器{iid}"
            insts.append({"id": iid, "name": nm})
        if not insts:
            # 没有活跃仪器上跑的，不作为必做项
            continue

        # 找最近一次 plan（按 test_item.name 匹配）
        last_plan = latest_plan_by_name.get((ti.name or "").strip())
        if last_plan:
            last_status = last_plan.status or "draft"
            last_label = f"{last_plan.year}年{'上半年' if last_plan.half == 1 else '下半年'}"
            # 下次应做：以上次比对日期 + 6 个月；无比对日期时以该半年度最后一天为基准
            base_date = _parse_date_str(last_plan.compared_at) or _half_year_end(last_plan.year, last_plan.half)
            next_due_date = _add_months(base_date, 6)
            next_due_label = _fmt_before(next_due_date)
            progress = {
                "last_plan": last_label,
                "last_status": "done" if last_status == "done" else "in_progress",
                "next_due": next_due_label,
            }
        else:
            progress = {"last_plan": "", "last_status": "never", "next_due": "请尽快安排"}

        out.append({
            "id": ti.id,
            "code": (ti.code or "").strip(),
            "name": (ti.name or "").strip(),
            "unit": (ti.unit or "").strip(),
            "instruments": insts,
            **progress,
        })
    return out


# ---------------------------------------------------------------------------
# 计算结果（5 水平）
# ---------------------------------------------------------------------------
def compute_data(db, plan, items, levels_map: dict):
    """items: list[InterlabItem]; levels_map: {item_id -> list[InterlabLevel]}。
    返回每个项目的 5 水平计算结果。

    语义（用户更正）：
    - 可比较系统 = 参比实验室 → ref1_*
    - 比较系统1 = 本实验室平台1 → our_value
    - 比较系统2 = 本实验室平台2 → ref2_*（可能有，亦可能空缺）
    偏倚均指「本实验室比较系统 vs 可比较系统(参比)」。
    """
    project_rows = []
    all_ok = True
    has_qual = False
    has_quan = False

    for it in items:
        kind = getattr(it, "kind", None) or "定量"
        levels = sorted(levels_map.get(it.id, []), key=lambda x: x.level_num)

        level_results = []
        if kind == "定性":
            has_qual = True
            for lv in levels:
                # 比较系统1（本实验室平台1）vs 可比较系统（参比）
                matched1, pn1 = _qualitative_eval(lv.our_value, lv.ref1_y1)
                has2 = bool((lv.ref2_y1 or "").strip())
                matched2, pn2 = (None, "") if not has2 else _qualitative_eval(lv.ref2_y1, lv.ref1_y1)
                acc1 = matched1
                acc2 = matched2 if has2 else None
                if acc1 is False:
                    all_ok = False
                if has2 and acc2 is False:
                    all_ok = False
                level_results.append({
                    "level": lv.level_num,
                    "ref": lv.ref1_y1,
                    "sys1": lv.our_value, "pn1": pn1, "match1": matched1, "accepted1": acc1,
                    "sys2": lv.ref2_y1 if has2 else "", "pn2": pn2 if has2 else "",
                    "match2": matched2 if has2 else None, "accepted2": acc2,
                    "has2": has2,
                })
        else:
            has_quan = True
            for lv in levels:
                ref_val = lv.ref1_mean or lv.ref1_y1 or ""
                of = _parse_float(lv.our_value)
                rf = _parse_float(ref_val)
                tf = _parse_float(it.te)
                abs_bias1 = (of - rf) if (of is not None and rf is not None) else None
                rel_bias1 = ((of - rf) / rf * 100.0) if (of is not None and rf not in (None, 0)) else None
                if it.mode == "absolute":
                    bias1 = abs_bias1
                    accepted1 = (abs_bias1 is not None and tf is not None and abs(abs_bias1) <= tf + 1e-9)
                else:
                    bias1 = rel_bias1
                    accepted1 = (rel_bias1 is not None and tf is not None and abs(rel_bias1) <= tf + 1e-9)
                if accepted1 is False:
                    all_ok = False

                r2 = _parse_float(lv.ref2_mean or lv.ref2_y1 or "")
                has2 = r2 is not None
                abs_bias2 = (r2 - rf) if (has2 and rf is not None) else None
                rel_bias2 = ((r2 - rf) / rf * 100.0) if (has2 and rf not in (None, 0)) else None
                if it.mode == "absolute":
                    bias2 = abs_bias2
                    accepted2 = (abs_bias2 is not None and tf is not None and abs(abs_bias2) <= tf + 1e-9)
                else:
                    bias2 = rel_bias2
                    accepted2 = (rel_bias2 is not None and tf is not None and abs(rel_bias2) <= tf + 1e-9)
                if has2 and accepted2 is False:
                    all_ok = False

                level_results.append({
                    "level": lv.level_num,
                    "ref_y1": lv.ref1_y1, "ref_y2": lv.ref1_y2, "ref_mean": lv.ref1_mean,
                    "sys1": lv.our_value,
                    "abs_bias1": round(abs_bias1, 3) if abs_bias1 is not None else None,
                    "rel_bias1": round(rel_bias1, 2) if rel_bias1 is not None else None,
                    "bias1": bias1, "accepted1": accepted1,
                    "sys2": (lv.ref2_mean or lv.ref2_y1) if has2 else "",
                    "abs_bias2": round(abs_bias2, 3) if abs_bias2 is not None else None,
                    "rel_bias2": round(rel_bias2, 2) if rel_bias2 is not None else None,
                    "bias2": bias2, "accepted2": accepted2,
                    "has2": has2,
                })

        proj_has2 = any(l.get("has2") for l in level_results)
        project_rows.append({
            "item": it.item, "unit": it.unit,
            "te": it.te, "mode": it.mode, "kind": kind, "note": it.note,
            "levels": level_results, "has2": proj_has2,
        })

    return {"projects": project_rows, "all_ok": all_ok, "has_qual": has_qual, "has_quan": has_quan}


def _qualitative_eval(our, ref):
    a, b = _norm_pn(our), _norm_pn(ref)
    if a is None or b is None:
        return None, f"{our or '—'} / {ref or '—'}"
    matched = (a == b)
    label = "阳性" if a == "positive" else "阴性"
    return matched, f"{label} / {label if matched else ('阴性' if a == 'positive' else '阳性')}"


# ---------------------------------------------------------------------------
# HTML 预览
# ---------------------------------------------------------------------------
def build_html(plan, data: dict, instrument_name: str, ref_lab: str):
    css = """
    <style>
      .rep { font-family:'Microsoft YaHei','SimSun',serif; color:#222; max-width:960px; margin:0 auto; padding:16px; }
      .rep h1 { text-align:center; font-size:20px; margin:6px 0; }
      .rep .sub { text-align:center; font-size:12px; color:#555; margin-bottom:10px; }
      .rep h2 { font-size:14px; margin:14px 0 6px; border-left:4px solid #1a365d; padding-left:8px; }
      .rep .note { font-size:12px; color:#666; margin:4px 0; }
      .rep table { border-collapse:collapse; width:100%; font-size:12px; margin:8px 0; }
      .rep th, .rep td { border:1px solid #444; padding:4px 6px; text-align:center; }
      .rep th { background:#eef2f8; }
      .rep .item { text-align:left; font-weight:bold; background:#f8fafe; }
      .rep .no { color:#c0392b; font-weight:700; }
      .rep .yes { color:#27ae60; font-weight:700; }
      .rep .concl { margin-top:6px; font-size:13px; }
      .rep .proj-sep { margin-top:14px; }
    </style>"""
    half = "上半年" if plan.half == 1 else "下半年"
    yn = plan.year or ""

    if data["has_qual"] and not data["has_quan"]:
        title = "定性项目室间比对结果记录及分析报告表"
    elif data["has_quan"] and not data["has_qual"]:
        title = "定量项目室间比对结果记录及分析报告表"
    else:
        title = "室间比对结果记录及分析报告表"

    html = [f'<div class="rep">{css}<h1>{title}</h1>']
    table_no = "BG-SM-CZ-019" if data["has_quan"] else "BG-SM-CZ-018"
    html.append(f'<div class="sub">表格编号：{table_no}　　民航总医院检验科生化免疫组　　生效日期：2026.1.1</div>')
    sys2 = getattr(plan, "compared_instrument2", "") or ""
    html.append("<h2>基本信息</h2>")
    html.append(f"<p>本实验室仪器（比较系统1）：{instrument_name or '　　　　'}　　参比实验室（可比较系统）：{ref_lab or '　　　　'}</p>")
    if sys2:
        html.append(f"<p>本实验室比较系统2：{sys2}</p>")
    html.append(f"<p>比对日期：{plan.compared_at or '　　　　'}　　操作者：{plan.operator or '　　　　'}　　审核者：{plan.reviewer or '　　　　'}</p>")

    # ---- 定量（BG-SM-CZ-019 结构：每项目一张 5 水平表） ----
    if data["has_quan"]:
        html.append("<h2>定量项目室间比对结果</h2>")
        html.append('<p class="note">注：参考WS/T415-2024《无室间质量评价时的临床检验质量评价》。<br>80%样本（4/5）的相对偏倚需低于允许总误差认为合格，允许总误差参照行标及国家卫健委室间质评要求，无相应要求的按照30%计算。</p>')
        for proj in data["projects"]:
            if proj["kind"] != "定量":
                continue
            has2 = proj.get("has2")
            html.append(f'<p style="margin:8px 0 2px;font-weight:bold;">项目：{proj["item"]}（单位：{proj["unit"]}，允许TE：{proj["te"]}{"%" if proj["mode"]=="relative" else ""}）</p>')
            head = ('<th>水平</th><th>参比值Y<br>(可比较系统)</th>'
                    '<th>比较系统1值<br>(本实验室)</th><th>比较系统1偏倚</th>')
            if has2:
                head += '<th>比较系统2值<br>(本实验室)</th><th>比较系统2偏倚</th>'
            head += '<th>是否合格</th>'
            html.append(f'<table><thead><tr>{head}</tr></thead><tbody>')
            ok_count = 0
            for lv in proj["levels"]:
                if proj["mode"] == "absolute":
                    b1 = f"{lv['abs_bias1']}" if lv.get("abs_bias1") is not None else "—"
                    b2 = f"{lv['abs_bias2']}" if (has2 and lv.get("abs_bias2") is not None) else "—"
                else:
                    b1 = f"{lv['rel_bias1']}%" if lv.get("rel_bias1") is not None else "—"
                    b2 = f"{lv['rel_bias2']}%" if (has2 and lv.get("rel_bias2") is not None) else "—"
                lev_ok = lv["accepted1"] and (lv["accepted2"] if has2 else True)
                if lev_ok:
                    ok_count += 1
                acc_cls = "yes" if lev_ok else ("no" if (lv["accepted1"] is False or (has2 and lv["accepted2"] is False)) else "")
                acc_s = "合格" if lev_ok else "不合格"
                row = (f'<tr><td>{lv["level"]}</td><td>{lv["ref_y1"]}</td>'
                       f'<td>{lv["sys1"]}</td><td>{b1}</td>')
                if has2:
                    row += f'<td>{lv["sys2"]}</td><td>{b2}</td>'
                row += f'<td class="{acc_cls}">{acc_s}</td></tr>'
                html.append(row)
            html.append("</tbody></table>")
            html.append(f'<p class="concl">5个水平中合格{ok_count}个（4/5即通过），项目{proj["item"]}室间比对一致性'
                        f'{"可接受" if ok_count >= 4 else "不可接受"}。</p>')

    # ---- 定性（BG-SM-CZ-018 结构） ----
    if data["has_qual"]:
        html.append("<h2>定性项目室间比对结果</h2>")
        for proj in data["projects"]:
            if proj["kind"] != "定性":
                continue
            has2 = proj.get("has2")
            html.append(f'<p style="margin:8px 0 2px;font-weight:bold;">项目：{proj["item"]}</p>')
            head = ('<th>水平</th><th>参比结果<br>(可比较系统)</th>'
                    '<th>比较系统1<br>(本实验室)</th><th>比较系统1符合</th>')
            if has2:
                head += '<th>比较系统2<br>(本实验室)</th><th>比较系统2符合</th>'
            head += '<th>是否合格</th>'
            html.append(f'<table><thead><tr>{head}</tr></thead><tbody>')
            matched = 0
            for lv in proj["levels"]:
                m_s1 = {True: "是", False: "否", None: "-"}[lv.get("match1")]
                lev_ok = lv["accepted1"] and (lv["accepted2"] if has2 else True)
                if lev_ok:
                    matched += 1
                acc_cls = "yes" if lev_ok else ("no" if (lv["accepted1"] is False or (has2 and lv["accepted2"] is False)) else "")
                acc_s = "合格" if lev_ok else "不合格"
                row = (f'<tr><td>{lv["level"]}</td><td>{lv["ref"]}</td>'
                       f'<td>{lv["sys1"]}</td><td>{m_s1}</td>')
                if has2:
                    m_s2 = {True: "是", False: "否", None: "-"}[lv.get("match2")]
                    row += f'<td>{lv["sys2"]}</td><td>{m_s2}</td>'
                row += f'<td class="{acc_cls}">{acc_s}</td></tr>'
                html.append(row)
            html.append("</tbody></table>")
            rate = round(matched / max(len(proj["levels"]), 1) * 100, 1)
            html.append(f'<p class="concl">5个水平符合{matched}/{len(proj["levels"])}，符合率{rate}%。</p>')

    html.append(f'<h2>结果分析</h2><p>{plan.summary or "各项目室间比对结果均在允许范围内，比对可接受。"}</p>')
    html.append(f'<h2>处理方案（如不合格）</h2><p>{plan.handle_plan or "无"}</p>')
    concl = plan.conclusion or ("可接受" if data["all_ok"] else "不可接受")
    html.append(f'<p>结论：<b>{concl}</b></p>')
    html.append(f'<div class="foot" style="margin-top:18px;">操作者：{plan.operator or "　　　　"}　　审核者：{plan.reviewer or "　　　　"}　　日期：{plan.compared_at or "　　　　"}</div>')
    html.append("</div>")
    return "".join(html)


# ---------------------------------------------------------------------------
# DOCX 生成
# ---------------------------------------------------------------------------
def _set_cell_font(cell, size=10.5, bold=False, align="center"):
    for p in cell.paragraphs:
        p.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.name = "SimSun"
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _fill(cell, text, size=10.5, bold=False, align="center", color=None):
    cell.text = str(text)
    _set_cell_font(cell, size, bold, align)
    if color:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = color


def _heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    return p


def _add_basic_info(doc, instrument_name, ref_lab, plan):
    sys2 = getattr(plan, "compared_instrument2", "") or ""
    _heading(doc, "基本信息")
    lines = [
        f"本实验室仪器（比较系统1）：{instrument_name or '　　　　'}　　参比实验室（可比较系统）：{ref_lab or '　　　　'}",
    ]
    if sys2:
        lines.append(f"本实验室比较系统2：{sys2}")
    lines.append(f"比对日期：{plan.compared_at or '　　　　'}　　操作者：{plan.operator or '　　　　'}　　审核者：{plan.reviewer or '　　　　'}")
    for line in lines:
        p = doc.add_paragraph()
        rr = p.add_run(line)
        rr.font.size = Pt(11)
        rr.font.name = "SimSun"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _add_footer(doc, plan):
    p = doc.add_paragraph()
    rf = p.add_run(f"操作者：{plan.operator or '　　　　'}　　审核者：{plan.reviewer or '　　　　'}　　日期：{plan.compared_at or '　　　　'}")
    rf.font.size = Pt(11)
    rf.font.name = "SimSun"
    rf._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def build_docx(db, plan, data: dict, out_path: str, instrument_name: str, ref_lab: str):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    year = plan.year or ""
    half = "上半年" if plan.half == 1 else "下半年"
    has_q = data["has_qual"]
    has_n = data["has_quan"]

    if has_q and not has_n:
        title = "定性项目室间比对结果记录及分析报告表"
    elif has_n and not has_q:
        title = "定量项目室间比对结果记录及分析报告表"
    else:
        title = "室间比对结果记录及分析报告表"

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title)
    r.font.size = Pt(16); r.font.bold = True; r.font.name = "SimSun"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_no = "BG-SM-CZ-019" if has_n else "BG-SM-CZ-018"
    rs = sub.add_run(f"表格编号：{table_no}　　民航总医院检验科生化免疫组　　生效日期：2026.1.1")
    rs.font.size = Pt(10.5); rs.font.name = "SimSun"
    rs._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    _add_basic_info(doc, instrument_name, ref_lab, plan)

    # ---- 定量表 ----
    if has_n:
        _heading(doc, "定量项目室间比对结果")
        pn = doc.add_paragraph()
        rr = pn.add_run("注：参考WS/T415-2024《无室间质量评价时的临床检验质量评价》。")
        rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66); rr.font.name = "SimSun"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        pn2 = doc.add_paragraph()
        rr2 = pn2.add_run("80%样本（4/5）的相对偏倚需低于允许总误差认为合格，允许总误差参照行标及国家卫健委室间质评要求，无相应要求的按照30%计算。")
        rr2.font.size = Pt(9.5); rr2.font.color.rgb = RGBColor(0x66, 0x66, 0x66); rr2.font.name = "SimSun"
        rr2._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        for proj in data["projects"]:
            if proj["kind"] != "定量":
                continue
            # 项目标题行
            p = doc.add_paragraph()
            pr = p.add_run(f"项目：{proj['item']}（单位：{proj['unit']}，允许TE：{proj['te']}{'%' if proj['mode']=='relative' else ''}）")
            pr.font.size = Pt(10.5); pr.font.bold = True; pr.font.name = "SimSun"
            pr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

            has2 = proj.get("has2")
            headers = ["水平", "参比值Y\n(可比较系统)", "比较系统1值\n(本实验室)", "比较系统1偏倚"]
            if has2:
                headers += ["比较系统2值\n(本实验室)", "比较系统2偏倚"]
            headers += ["是否合格"]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            for i, ht in enumerate(headers):
                _fill(hdr[i], ht, bold=True)

            ok_count = 0
            for lv in proj["levels"]:
                if proj["mode"] == "absolute":
                    b1 = f"{lv['abs_bias1']}" if lv.get("abs_bias1") is not None else "—"
                    b2 = f"{lv['abs_bias2']}" if (has2 and lv.get("abs_bias2") is not None) else "—"
                else:
                    b1 = f"{lv['rel_bias1']}%" if lv.get("rel_bias1") is not None else "—"
                    b2 = f"{lv['rel_bias2']}%" if (has2 and lv.get("rel_bias2") is not None) else "—"
                lev_ok = lv["accepted1"] and (lv["accepted2"] if has2 else True)
                if lev_ok:
                    ok_count += 1
                acc_s = "合格" if lev_ok else "不合格"
                cells = t.add_row().cells
                _fill(cells[0], str(lv["level"]))
                _fill(cells[1], lv["ref_y1"])
                _fill(cells[2], lv["sys1"])
                _fill(cells[3], b1)
                if has2:
                    _fill(cells[4], lv["sys2"])
                    _fill(cells[5], b2)
                col_idx = 5 if has2 else 4
                col = RGBColor(0x27, 0xae, 0x60) if lev_ok else RGBColor(0xc0, 0x39, 0x2b)
                _fill(cells[col_idx], acc_s, color=col, bold=True)

            p = doc.add_paragraph()
            pr = p.add_run(f"结论：5个水平中合格{ok_count}个（4/5即通过），项目{proj['item']}室间比对一致性{'可接受' if ok_count >= 4 else '不可接受'}。")
            pr.font.size = Pt(10.5); pr.font.name = "SimSun"
            pr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # ---- 定性表 ----
    if has_q:
        _heading(doc, "定性项目室间比对结果")
        for proj in data["projects"]:
            if proj["kind"] != "定性":
                continue
            p = doc.add_paragraph()
            pr = p.add_run(f"项目：{proj['item']}")
            pr.font.size = Pt(10.5); pr.font.bold = True; pr.font.name = "SimSun"
            pr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

            has2 = proj.get("has2")
            headers = ["水平", "参比结果\n(可比较系统)", "比较系统1\n(本实验室)", "比较系统1符合"]
            if has2:
                headers += ["比较系统2\n(本实验室)", "比较系统2符合"]
            headers += ["是否合格"]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            for i, ht in enumerate(headers):
                _fill(hdr[i], ht, bold=True)

            matched = 0
            for lv in proj["levels"]:
                m_s1 = {True: "是", False: "否", None: "-"}[lv.get("match1")]
                lev_ok = lv["accepted1"] and (lv["accepted2"] if has2 else True)
                if lev_ok:
                    matched += 1
                acc_s = "合格" if lev_ok else "不合格"
                cells = t.add_row().cells
                _fill(cells[0], str(lv["level"]))
                _fill(cells[1], lv["ref"])
                _fill(cells[2], lv["sys1"])
                _fill(cells[3], m_s1)
                if has2:
                    m_s2 = {True: "是", False: "否", None: "-"}[lv.get("match2")]
                    _fill(cells[4], lv["sys2"])
                    _fill(cells[5], m_s2)
                col_idx = 5 if has2 else 4
                col = RGBColor(0x27, 0xae, 0x60) if lev_ok else RGBColor(0xc0, 0x39, 0x2b)
                _fill(cells[col_idx], acc_s, color=col, bold=True)

            rate = round(matched / max(len(proj["levels"]), 1) * 100, 1)
            p = doc.add_paragraph()
            pr = p.add_run(f"结论：5个水平符合{matched}/{len(proj['levels'])}，符合率{rate}%。")
            pr.font.size = Pt(10.5); pr.font.name = "SimSun"
            pr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    _heading(doc, "结果分析")
    p = doc.add_paragraph()
    rr = p.add_run(plan.summary or "各项目室间比对结果均在允许范围内，比对可接受。")
    rr.font.size = Pt(11); rr.font.name = "SimSun"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    _heading(doc, "处理方案（如不合格）")
    p = doc.add_paragraph()
    rr = p.add_run(plan.handle_plan or "无")
    rr.font.size = Pt(11); rr.font.name = "SimSun"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    concl = plan.conclusion or ("可接受" if data["all_ok"] else "不可接受")
    p = doc.add_paragraph()
    rr = p.add_run(f"结论：{concl}")
    rr.font.size = Pt(11); rr.font.bold = True; rr.font.name = "SimSun"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    _add_footer(doc, plan)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
