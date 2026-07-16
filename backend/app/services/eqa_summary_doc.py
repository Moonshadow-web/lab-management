"""室间质评半年/年度总结 Word 生成（基于真实 BG-SM-CZ-020 模板填充）。

做法：复制 data/uploads/docs/BG-SM-CZ-020-生化免疫组室间质评总结报告表.docx
作为底版，仅按字段填入学统计数字与文字，保留模板原有字体、合并单元格、
边框与签字区版式；不从头画线框表格。
"""
from docx import Document
from pathlib import Path
import shutil


def _period_label(half: int) -> str:
    return {1: "上半年（1-6月）", 2: "下半年（7-12月）", 0: "全年"}.get(half, "全年")


# 质评部门（org）显示名映射
_ORG_DISPLAY = {"卫健委": "国家卫健委临检中心", "北京市": "北京市临检中心"}


def _find_template() -> Path:
    """定位真实 BG-SM-CZ-020 模板（项目根/data/uploads/docs/...）。"""
    base = Path(__file__).resolve().parents[3]
    cand = base / "data" / "uploads" / "docs" / "BG-SM-CZ-020-生化免疫组室间质评总结报告表.docx"
    if cand.exists():
        return cand
    hits = list((base / "data").rglob("*BG-SM-CZ-020*.docx"))
    if hits:
        return hits[0]
    raise FileNotFoundError("未找到 BG-SM-CZ-020 室间质评总结模板")


def _fill_by_text(doc, substr: str, new_text: str):
    """把所有包含 substr 的单元格整格替换为 new_text（自动兼容合并/独立单元格）。"""
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                if substr in cell.text:
                    cell.text = new_text


def _set_cell_multiline(cell, lines):
    """清空单元格并按多行写入（保留模板段落字体）。"""
    # 清空所有段落
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        p.add_run(line)


def build_eqa_summary_doc(year, half, stats, summary_text, out_path, category=None, department=None):
    """基于真实 BG-SM-CZ-020 模板生成总结 docx 并保存到 out_path。

    stats: _compute_summary_by_category 返回的结构（含 categories / total，已去重）。
    category: 专业组（生化+凝血 / 免疫）；department: 质评部门（卫健委 / 北京市）。
    """
    dept_key = (department or "").strip()
    dept_display = _ORG_DISPLAY.get(dept_key, dept_key or "—")
    cat = (category or "").strip()
    period = _period_label(half)

    t = stats.get("total", {})
    items_total = t.get("items_total", 0)
    qualified = t.get("qualified", 0)
    unqualified = t.get("unqualified", 0)
    not_evaluated = t.get("not_evaluated", 0)
    rate = t.get("qualify_rate")
    ev = t.get("items_evaluated", 0)

    # 1. 复制真实模板为底版
    tmpl = _find_template()
    shutil.copy(str(tmpl), str(out_path))
    doc = Document(str(out_path))

    # 2. 年份 / 质评部门（模板 R1）
    _fill_by_text(doc, "年份：", f"年份：{year}年（{period}）")
    org_text = f"质评部门：{dept_display}（{cat}组）" if cat else f"质评部门：{dept_display}"
    _fill_by_text(doc, "质评部门：", org_text)

    # 3. 测定项目概览（模板 R2，整行同内容）
    overview = (
        f"测定项目：共 1 组别（{cat}）， {items_total} 项。"
        f"合格 {qualified} 项，不合格 {unqualified} 项。"
    )
    _fill_by_text(doc, "测定项目：共", overview)

    # 4. 不合格项目明细（模板 R3 表头后的 3 行）
    unq = []
    for c in stats.get("categories", []):
        for it in c.get("unqualified_list", []):
            unq.append(it)
    tb = doc.tables[0]
    hdr_idx = None
    for i, row in enumerate(tb.rows):
        if "不合格项目" in row.cells[0].text and "上报结果" in row.cells[1].text:
            hdr_idx = i
            break
    if hdr_idx is not None:
        for k in range(3):
            r = tb.rows[hdr_idx + 1 + k]
            if k < len(unq):
                it = unq[k]
                r.cells[0].text = it.get("item", "")
                r.cells[1].text = it.get("result") or it.get("score") or ""
                r.cells[2].text = "见官方成绩报告允许范围"
                r.cells[3].text = "（待填：原因分析及纠正措施）"
                r.cells[4].text = "不合格项目未对患者标本的结果造成影响"
            # 无更多不合格项则保持模板空白

    # 5. 合格项目分析（模板 R7，整行同内容，多行）
    first_line = (
        f"{year}年{dept_display}（{cat}组）室间质评细项合格率 "
        f"{rate}% （合格 {qualified} / 评价 {ev} 项）"
    )
    if not_evaluated:
        first_line += f"；另有 {not_evaluated} 项不予评价（单列，不计入合格率）。"
    lines = ["合格项目分析：", first_line]
    if summary_text and summary_text.strip():
        for ln in summary_text.strip().split("\n"):
            if ln.strip():
                lines.append(ln.strip())
    for row in tb.rows:
        for cell in row.cells:
            if "合格项目分析：" in cell.text:
                _set_cell_multiline(cell, lines)

    doc.save(str(out_path))
    return str(out_path)
