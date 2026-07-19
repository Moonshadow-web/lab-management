import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))  # backend/
from app.core.database import SessionLocal
from app.models.comparison import ComparisonPlan, ComparisonGroup
from app.services import comparison_report as svc

db = SessionLocal()
# 取一个定量计划
p = db.query(ComparisonPlan).filter_by(id=9).first()
g = db.get(ComparisonGroup, p.group_id)
data = svc.compute_data(db, g, p)

# 强制把水平1第一行改成绝对偏倚，并给一个偏倚值
blk = data["matrix"][0]
row = blk["rows"][0]
row["mode"] = "absolute"
for cid, c in row["insts"].items():
    if not c.get("masked") and not c.get("is_ref"):
        c["bias"] = -0.03
        c["accepted"] = True

html = svc.build_html(g, p, data)
assert "（绝对偏倚）" in html, "HTML 未包含 (绝对偏倚) 标记"
# 列头应标注相对偏倚（绝对偏倚行单元格已单独标注）
assert "相对偏倚" in html, "HTML 列头未包含 (相对偏倚)"
print("HTML 含 (绝对偏倚)单元格:", "（绝对偏倚）" in html, "| 含 (相对偏倚)列头:", "相对偏倚" in html)

out = os.path.join(os.path.dirname(__file__), "..", "tmp_abs_test.docx")
svc.build_docx(db, g, p, data, out)
from docx import Document
doc = Document(out)
text = "\n".join(par.text for par in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for cell in r.cells:
            text += "\n" + cell.text
assert "（绝对偏倚）" in text, "DOCX 未包含 (绝对偏倚) 标记"
assert "相对偏倚" in text, "DOCX 列头未包含 (相对偏倚)"
print("DOCX 含 (绝对偏倚)单元格:", "（绝对偏倚）" in text, "| 含 (相对偏倚)列头:", "相对偏倚" in text)
os.remove(out)
print("TEST OK")
db.close()
