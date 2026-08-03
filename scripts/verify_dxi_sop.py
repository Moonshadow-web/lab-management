#!/usr/bin/env python
# 校验 DXI 改写结果：旧 SOP-1004/1006 应清零；出现的 SOP 应全部等于该项目档案机器集合。
import os, re, urllib.request, urllib.parse, json
from docx import Document

FOLDER = r"C:/Users/81526/Desktop/待办/DXI"
BASE = "http://lab-management-282724-9-1408547492.sh.run.tcloudbase.com/api/v1"

def req(method, path, data=None, tok=""):
    h = {"Accept": "application/json"}
    if tok:
        h["Authorization"] = f"Bearer {tok}"
    body = None
    if data is not None:
        if method == "POST" and path == "/auth/login":
            body = urllib.parse.urlencode(data).encode()
            h["Content-Type"] = "application/x-www-form-urlencoded"
        else:
            body = json.dumps(data, ensure_ascii=False).encode()
            h["Content-Type"] = "application/json"
    r = urllib.request.Request(BASE + path, data=body, headers=h, method=method)
    try:
        resp = urllib.request.urlopen(r, timeout=120)
        return resp.status, json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        return e.code, e.read().decode()[:200]

st, l = req("POST", "/auth/login", {"username": "jinzizheng", "password": "Jzz6827556"})
tok = l["access_token"]
st, r = req("GET", "/test-items?page=1&page_size=5000", tok=tok)
id_to_name = {it["id"]: (it.get("name") or "").strip() for it in r.get("items", [])}

INSTS = {69: ("1号机", "2004"), 70: ("2号机", "2005"), 71: ("3号机", "2006"),
         72: ("4号机", "2007"), 73: ("急诊", "2008"), 74: ("唐筛", "2009")}
pm = {}
for iid, (fam, sop) in INSTS.items():
    st, r2 = req("GET", f"/instruments/{iid}/test-items", tok=tok)
    for it in r2:
        nm = id_to_name.get(it.get("id"))
        if nm:
            pm.setdefault(nm, set()).add(sop)

_alias = {
    "总β-人绒毛膜促性腺激素β-HCG": "β人绒毛膜促性腺激素",
    "β-HCG": "β人绒毛膜促性腺激素",
    "雌二醇E2": "雌二醇",
    "未结合雌三醇u-E3": "非结合型雌三醇",
    "未结合雌三醇E3": "非结合型雌三醇",
    "总25-羟维生素D": "维生素D",
    "25-羟维生素D": "维生素D",
    "全段甲状旁腺激素iPTH": "甲状旁腺激素",
    "甲状旁腺激素iPTH": "甲状旁腺激素",
    "超敏肌钙蛋白I": "超敏肌钙蛋白Ⅰ",
    "超敏肌钙蛋白": "超敏肌钙蛋白Ⅰ",
    "可溶性转铁蛋白受体sTfR": "可溶性转铁蛋白受体",
    "sTfR": "可溶性转铁蛋白受体",
    "肌酸激酶同工酶CK-MB": "肌酸激酶同工酶",
    "CK-MB": "肌酸激酶同工酶",
    "甲胎蛋白AFP": "甲胎蛋白",
    "AFP": "甲胎蛋白",
    "总前列腺特异性抗原TPSA": "总前列腺特异性抗原",
    "TPSA": "总前列腺特异性抗原",
    "游离前列腺特异性抗原TPSA": "游离前列腺特异性抗原",
    "游离前列腺特异性抗原fPSA": "游离前列腺特异性抗原",
    "抗缪勒管激素AMH": "抗缪勒管激素",
    "AMH": "抗缪勒管激素",
    "甲状腺球蛋白抗体TgAb": "抗甲状腺球蛋白抗体",
    "TgAb": "抗甲状腺球蛋白抗体",
    "抗甲状腺过氧化物酶抗体TPOAb": "抗甲状腺过氧化物酶抗体",
    "TPOAb": "抗甲状腺过氧化物酶抗体",
    "促红细胞生成素EPO": "促红细胞生成素",
    "EPO": "促红细胞生成素",
    "B型钠尿肽BNP": "B型钠尿肽",
    "BNP": "B型钠尿肽",
    "肌红蛋白MYO": "肌红蛋白",
    "MYO": "肌红蛋白",
    "叶酸FLOW": "叶酸",
    "内因子抗体IFAb": "抗内因子抗体",
    "IFAb": "抗内因子抗体",
    "总甲状腺素TT4": "甲状腺素",
    "总三碘甲状腺原氨酸TT3": "三碘甲状腺原氨酸",
    "游离甲状腺素FT4": "游离甲状腺素",
    "游离三碘甲状腺原氨酸FT3": "游离三碘甲状腺原氨酸",
    "促甲状腺激素TSH": "促甲状腺激素",
    "白介素6": "白介素-6",
    "血清降钙素原": "降钙素原",
    "血浆降钙素原": "降钙素原",
    "血清血浆降钙素原": "降钙素原",
    "血浆白介素6": "白介素-6",
    "血清白介素6": "白介素-6",
    "血清血浆白介素6": "白介素-6",
    "血浆铁蛋白": "铁蛋白",
    "血清铁蛋白": "铁蛋白",
    "血清血浆铁蛋白": "铁蛋白",
    "血浆可溶性转铁蛋白受体": "可溶性转铁蛋白受体",
    "血浆内因子抗体": "抗内因子抗体",
    "血浆叶酸": "叶酸",
    "血清血浆叶酸": "叶酸",
    "血浆维生素B12": "维生素B12",
    "血清血浆维生素B12": "维生素B12",
    "血浆B型钠尿肽": "B型钠尿肽",
    "血清血浆B型钠尿肽": "B型钠尿肽",
    "血浆肌红蛋白": "肌红蛋白",
    "血清血浆肌红蛋白": "肌红蛋白",
    "血浆超敏肌钙蛋白I": "超敏肌钙蛋白Ⅰ",
    "血清血浆超敏肌钙蛋白I": "超敏肌钙蛋白Ⅰ",
    "血浆超敏肌钙蛋白": "超敏肌钙蛋白Ⅰ",
    "血清血浆超敏肌钙蛋白": "超敏肌钙蛋白Ⅰ",
    "血浆肌酸激酶同工酶": "肌酸激酶同工酶",
    "血清血浆肌酸激酶同工酶": "肌酸激酶同工酶",
    "血清降钙素": "降钙素",
    "血浆降钙素": "降钙素",
}

def parse_project(fn):
    fn = fn.rsplit(".", 1)[0]
    fn = re.sub(r"\s+NA$", "", fn).strip()
    fn = re.sub(r"^SM-SOP-\d+[\s\-]*", "", fn)
    fn = re.sub(r"^贝克曼DXI800检测系统", "", fn)
    fn = re.sub(r"^(?:化学发光法|电化学发光法|IGRA法|免疫分析|酶联免疫吸附法)\s*", "", fn)
    fn = re.sub(r"^(?:血清|血浆|尿液|全血|胸腹水|脑脊液)+", "", fn)
    fn = re.sub(r"\s+沃文特$", "", fn).strip()
    fn = re.sub(r"FLOW$", "", fn).strip()
    for _ in range(4):
        nw = re.sub(r"-+$", "", fn).strip()
        nw = re.sub(r"(?:测定标准操作程序|标准测定程序|标准操作程序|操作程序|测定)$", "", nw).strip()
        if nw == fn:
            break
        fn = nw
    return fn.strip()

SOPREF = re.compile(r"SOP-(\d{4}(?:/\d{4})*)")
OLD = {"1004", "1006"}
# 非 DXI 仪器家族的 SOP 编号（如 1002=AU5800 生化仪），不在本任务改写范围，仅作信息提示
NON_DXI_SOP = {"1002", "1004", "1006"}
DXI_SOPS = {"2004", "2005", "2006", "2007", "2008", "2009"}
problems = 0
checked = 0
for fn in sorted(os.listdir(FOLDER)):
    if not fn.lower().endswith(".docx") or fn.startswith("~$"):
        continue
    proj = parse_project(fn)
    name = _alias.get(proj, proj)
    exp = pm.get(name)
    doc = Document(os.path.join(FOLDER, fn))
    txt = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    txt += "\n" + p.text
    found = set()
    for m in SOPREF.findall(txt):
        found.update(m.split("/"))
    checked += 1
    if exp is None:
        print(f"[NOMATCH] {fn[:44]} project={proj}")
        problems += 1
        continue
    if found & OLD:
        print(f"[OLD-REMAIN] {fn[:44]} leftover={found & OLD}")
        problems += 1
    # 只校验 DXI 机器集合(2004-2009)与档案一致；非 DXI 编号(如1002=AU5800)仅提示
    dxi_found = found & DXI_SOPS
    other = found - DXI_SOPS - NON_DXI_SOP
    if other:
        print(f"[UNEXPECTED-SOP] {fn[:44]} non-DXI/unknown={sorted(other)}")
        problems += 1
    if dxi_found != set(exp):
        print(f"[MISMATCH] {fn[:44]} dxi_found={sorted(dxi_found)} expected={sorted(exp)}")
        problems += 1
    info_others = found & NON_DXI_SOP
    if info_others and dxi_found == set(exp):
        print(f"[INFO] {fn[:44]} 含非DXI引用(未改动)={sorted(info_others)}，DXI引用正确")
print(f"\nChecked={checked}  problems={problems}")
