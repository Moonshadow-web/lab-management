#!/usr/bin/env python
# 校验安图 A2000->A6200 改写结果：
#  - 旧串清零：AutolumoA2000(忽略大小写a2000形式) / AutoLumo A2000PLUS / SOP-1019 / 安图A2000PLUS
#  - 新串到位：AutoLumoA6200(标题) / AutoLumo A6200(正文) / SOP-2010 / 安图A6200
#  - SOP-2010 引用格式：详见编号MHZYY-JYK-SM-SOP-2010《安图A6200免疫分析仪操作作业指导书》
import os, re
from docx import Document

FOLDER = r"C:/Users/81526/Desktop/待办/安图"

def blob_of(fn):
    doc = Document(os.path.join(FOLDER, fn))
    b = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    b += "\n" + p.text
    return b, doc

problems = 0
checked = 0
old_a2000 = re.compile(r"(?i)autolumo\s*a\s*2000")          # 任何 Autolumo+2000(旧)
old_plus = re.compile(r"AutoLumo\s*A2000\s*PLUS")           # 正文旧
new_title = "AutoLumoA6200"
new_body = "AutoLumo A6200"
for fn in sorted(os.listdir(FOLDER)):
    if not fn.lower().endswith(".docx") or fn.startswith("~$"):
        continue
    b, doc = blob_of(fn)
    checked += 1
    o1 = len(old_a2000.findall(b))
    o2 = len(old_plus.findall(b))
    o3 = b.count("SOP-1019")
    o4 = b.count("安图A2000PLUS")
    n1 = b.count(new_title)
    n2 = b.count(new_body)
    n3 = b.count("SOP-2010")
    n4 = b.count("安图A6200")
    # SOP-2010 引用格式校验
    sop_ok = True
    sop_refs = re.findall(r"详见编号MHZYY-JYK-SM-SOP-2010《安图A6200免疫分析仪操作作业指导书》", b)
    # 若含 SOP-2010，必须格式完全正确；且不能出现 SOP-2010 的其他(错误)写法
    bad_sop = re.findall(r"SOP-2010《[^》]*》", b)
    bad_sop = [x for x in bad_sop if x != "SOP-2010《安图A6200免疫分析仪操作作业指导书》"]
    if n3:
        if not sop_refs or bad_sop:
            sop_ok = False
    issues = []
    if o1: issues.append(f"旧AutolumoA2000={o1}")
    if o2: issues.append(f"旧A2000PLUS={o2}")
    if o3: issues.append(f"旧SOP-1019={o3}")
    if o4: issues.append(f"旧安图A2000PLUS={o4}")
    if n1 == 0: issues.append("新标题AutoLumoA6200=0")
    if n2 == 0: issues.append("新正文AutoLumo A6200=0")
    if bad_sop: issues.append(f"SOP-2010错误写法={bad_sop[:1]}")
    if issues:
        problems += 1
        print(f"[PROBLEM] {fn[:46]}: {'; '.join(issues)}")
    else:
        tag = f"SOP-2010x{len(sop_refs)}" if n3 else "无SOP引用"
        print(f"[OK] {fn[:40]:42} 标题{new_title}x{n1} 正文{new_body}x{n2} {tag}")

print(f"\nChecked={checked}  problems={problems}")
# 全局汇总
allb = "\n".join(blob_of(f)[0] for f in sorted(os.listdir(FOLDER)) if f.lower().endswith(".docx") and not f.startswith("~$"))
print(f"全局旧串: AutolumoA2000类={len(old_a2000.findall(allb))} A2000PLUS={len(old_plus.findall(allb))} SOP-1019={allb.count('SOP-1019')} 安图A2000PLUS={allb.count('安图A2000PLUS')}")
print(f"全局新串: AutoLumoA6200={allb.count(new_title)} AutoLumo A6200={allb.count(new_body)} SOP-2010={allb.count('SOP-2010')} 安图A6200={allb.count('安图A6200')}")
