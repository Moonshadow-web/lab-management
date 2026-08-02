#!/usr/bin/env python
# 回退+修正 DXI 文档：
#  ① 项目标题误改回退：「测定操作作业指导书」→「测定标准操作程序」
#  ② 仪器SOP引用修正：「化学发光仪标准操作作业指导书」→「化学发光免疫分析仪操作作业指导书」
# 两者无重叠子串，顺序无关；run 级保留格式。
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn
import copy as _copy

FOLDER = r"C:/Users/81526/Desktop/待办/DXI"
BACKUP = r"C:/Users/81526/Desktop/待办/DXI_BACKUP_RB"
PAIRS = [
    ("测定操作作业指导书", "测定标准操作程序"),
    ("化学发光仪标准操作作业指导书", "化学发光免疫分析仪操作作业指导书"),
]
PATTERNS = [(re.compile(re.escape(o)), n) for o, n in PAIRS]

def copy_run_style(dst, src):
    f = src.font; df = dst.font
    if f.name: df.name = f.name
    if f.size: df.size = f.size
    if f.bold is not None: df.bold = f.bold
    if f.italic is not None: df.italic = f.italic
    if f.underline is not None: df.underline = f.underline
    if f.color and f.color.rgb is not None: df.color.rgb = f.color.rgb
    rpr = src._element.find(qn('w:rPr'))
    if rpr is not None: dst._element.append(_copy.deepcopy(rpr))

def replace_paragraph(par, pattern, make_repl):
    runs = par.runs
    if not runs: return False
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches: return False
    char_to_run = []
    for ri, r in enumerate(runs):
        char_to_run.extend([ri] * len(r.text))
    orig_runs = list(runs)
    segments = []; last = 0
    for m in matches:
        s, e = m.span()
        i = last
        while i < s:
            ri = char_to_run[i]; seg_start = i
            while i < s and char_to_run[i] == ri: i += 1
            segments.append((full[seg_start:i], ri))
        segments.append((make_repl(m), char_to_run[s]))
        last = e
    i = last
    while i < len(full):
        ri = char_to_run[i]; seg_start = i
        while i < len(full) and char_to_run[i] == ri: i += 1
        segments.append((full[seg_start:i], ri))
    for r in list(par.runs): r._element.getparent().remove(r._element)
    for text, ri in segments:
        if text == "": continue
        nr = par.add_run(text); copy_run_style(nr, orig_runs[ri])
    return True

def process(paras):
    c = 0
    for p in paras:
        for pat, new in PATTERNS:
            if replace_paragraph(p, pat, lambda m, n=new: n): c += 1
    return c

ONLY = os.environ.get("ONLY")
APPLY = os.environ.get("APPLY") == "1"
os.makedirs(BACKUP, exist_ok=True)
files = sorted(f for f in os.listdir(FOLDER) if f.lower().endswith(".docx") and not f.startswith("~$"))
if ONLY: files = [f for f in files if ONLY in f]
report = []
for fn in files:
    path = os.path.join(FOLDER, fn)
    doc = Document(path)
    changed = 0
    changed += process(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                changed += process(cell.paragraphs)
    if changed == 0:
        report.append((fn, 0, "SKIP"))
        continue
    if APPLY:
        shutil.copy2(path, os.path.join(BACKUP, fn))
        doc.save(path)
    report.append((fn, changed, "APPLY" if APPLY else "DRYRUN"))
print(f"{'APPLY' if APPLY else 'DRY-RUN'} files={len(report)}")
tot = sum(r[1] for r in report)
print(f"总替换段落/单元格处={tot}")
for fn, c, a in report:
    print(f"  {a:7} {c:3}  {fn[:54]}")
