"""解析 CNAS-AL02-07 附表3《医学实验室质量和能力认可准则和应用要求》自查表.docx
输出：clauses_export.json（clause_no/chapter/title/content/check_point）
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


def is_header_row(cells: list[str]) -> tuple[bool, str]:
    """判断是否为章节/条款标题行（所有非空单元格内容相同且含数字）。"""
    non_empty = [c for c in cells if c]
    if not non_empty:
        return False, ""
    first = non_empty[0]
    if not all(c == first for c in non_empty):
        return False, ""
    # 至少以数字开头
    if not re.match(r"^\d", first):
        return False, ""
    return True, first


def extract_clause_title(content: str) -> str:
    """从内容中抽取短标题；抽不到返回空。"""
    if not content:
        return ""
    # 1) 冒号/破折号前的短文本
    for sep in ("：", ":"):
        if sep in content:
            head = content.split(sep, 1)[0].strip()
            if 2 <= len(head) <= 40 and not _has_list_marker(head):
                return head.rstrip("：:")
    # 2) 空格/全角空格前的短文本
    parts = re.split(r"[ \s　]+", content, maxsplit=1)
    if parts and 2 <= len(parts[0]) <= 40 and not _has_list_marker(parts[0]):
        return parts[0].rstrip("：:")
    return ""


def _has_list_marker(s: str) -> bool:
    return bool(re.search(r"[a-z]\)|[a-z]）|\d+\)|\d+）", s))


CLAUSE_RE = re.compile(r"^(\d+(?:\.\d+){1,3})$")


def parse_docx(path: str) -> list[dict]:
    doc = Document(path)
    if not doc.tables:
        raise ValueError("文档中没有表格")
    tbl = doc.tables[0]

    clauses: list[dict] = []
    seen: set[str] = set()
    chapter = ""
    pending_title = ""
    pending_clause = ""

    for row in tbl.rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        if not cells:
            continue

        is_header, header_text = is_header_row(cells)
        if is_header:
            m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", header_text)
            if not m:
                continue
            num, title_text = m.group(1), m.group(2).strip()
            # 顶层章节（如 "4 总体要求"）作为 chapter，不生成条款
            if "." not in num:
                chapter = header_text
                pending_title = ""
                pending_clause = ""
                continue
            # 中间/条款标题行（如 "4.1 公正性" / "7.3.7 检验结果有效性的保证"）
            pending_title = title_text or num
            pending_clause = num
            # 同时生成一条空内容条款（保留分组标题，如 4.2 / 7.3.7）
            if num not in seen:
                seen.add(num)
                clauses.append({
                    "clause_no": num,
                    "chapter": chapter,
                    "title": pending_title,
                    "content": "",
                    "check_point": "",
                })
            continue

        first = cells[0]
        m = CLAUSE_RE.match(first)
        if not m:
            continue

        clause_no = m.group(1)
        content = cells[1] if len(cells) > 1 else ""
        check_point = cells[2] if len(cells) > 2 else ""

        # 应用要求列与认可准则列重复时，置空
        if check_point and check_point == content:
            check_point = ""

        # 标题来源优先级：内容自带标题 > 同组标题行
        title = extract_clause_title(content)
        if not title and pending_clause and (
            clause_no == pending_clause or clause_no.startswith(pending_clause + ".")
        ):
            title = pending_title
        if not title:
            title = clause_no

        # 避免内容里重复标题
        if title and title != clause_no and content.startswith(title):
            content = content[len(title):].lstrip(" ：:").strip()

        if clause_no not in seen:
            seen.add(clause_no)
            clauses.append({
                "clause_no": clause_no,
                "chapter": chapter,
                "title": title,
                "content": content,
                "check_point": check_point,
            })
        else:
            # 已存在则追加内容（处理偶发的同一条款多行）
            for c in clauses:
                if c["clause_no"] == clause_no:
                    if content:
                        c["content"] = (c["content"] + "\n" + content).strip()
                    if check_point:
                        c["check_point"] = (c["check_point"] + "\n" + check_point).strip()
                    break

    return clauses


def main():
    path = r"C:/Users/81526/Desktop/CNAS-AL02-07：20230801 附表3《医学实验室质量和能力认可准则和应用要求》自查表.docx"
    if len(sys.argv) > 1:
        path = sys.argv[1]
    out = Path(__file__).resolve().parent.parent / "clauses_export.json"

    clauses = parse_docx(path)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(clauses, f, ensure_ascii=False, indent=1)

    print("total clauses:", len(clauses))
    for c in clauses[:8]:
        print(f"  {c['clause_no']} | {c['chapter']} | {c['title']!r} | cLen={len(c['content'])} aLen={len(c['check_point'])}")
    print("...")
    for c in clauses:
        if c["clause_no"] in ("7.5", "7.3.7.2", "4.2", "7.3.7"):
            print(f"  SAMPLE {c['clause_no']} | {c['title']!r} | cLen={len(c['content'])} aLen={len(c['check_point'])}")


if __name__ == "__main__":
    main()
