#!/usr/bin/env python
# 替换 AU5800 系列项目 SOP 文档里对通用仪器 SOP-1002《AU5800生化分析仪标准操作程序》
# 的引用，改为该项目实际使用的机器（来自系统仪器档案 -> 使用本仪器的检验项目）。
import os, re, json, shutil, urllib.request, urllib.parse
from docx import Document
from docx.oxml.ns import qn

FOLDER = r"C:/Users/81526/Desktop/待办/AU5800"
BACKUP = r"C:/Users/81526/Desktop/待办/AU5800_BACKUP"
BASE="http://lab-management-282724-9-1408547492.sh.run.tcloudbase.com/api/v1"

# ---------- API ----------
def req(method,path,data=None,token=""):
    url=BASE+path; h={"Accept":"application/json"}
    if token: h["Authorization"]=f"Bearer {token}"
    body=None
    if data is not None:
        if method=="POST" and path=="/auth/login":
            body=urllib.parse.urlencode(data).encode(); h["Content-Type"]="application/x-www-form-urlencoded"
        else:
            body=json.dumps(data,ensure_ascii=False).encode(); h["Content-Type"]="application/json"
    r=urllib.request.Request(url,data=body,headers=h,method=method)
    try:
        resp=urllib.request.urlopen(r,timeout=120); return resp.status,json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        return e.code,e.read().decode()[:300]
st,l=req("POST","/auth/login",{"username":"jinzizheng","password":"Jzz6827556"}); tok=l["access_token"]

# 主表 id -> 规范名（仪器档案接口的 name 字段部分缺失，必须用 id 反查主表）
st,r=req("GET","/test-items?page=1&page_size=5000",token=tok)
id_to_name={}
for it in r.get("items",[]):
    id_to_name[it["id"]]=(it.get("name") or "").strip()

# 三台机器： id -> (档案名, SOP尾号, 机器显示名)
INSTS = {67:("AU58-1","2002","AU5821A"), 68:("AU58-2","2003","AU5821B"), 5:("AU5800","1005","AU5800急诊")}
proj_to_machines = {}
for iid,(fam,sop,name) in INSTS.items():
    st,r2=req("GET",f"/instruments/{iid}/test-items",token=tok)
    for it in r2:
        nm=id_to_name.get(it.get("id"))
        if nm: proj_to_machines.setdefault(nm,set()).add(fam)

# ---------- 文件名 -> 项目核心名 ----------
def parse_project(filename):
    fn = filename.rsplit(".",1)[0]
    fn = re.sub(r"\s+NA$","",fn).strip()
    fn = re.sub(r"^SM-SOP-\d+[\s\-]*(?:AU5800检测系统)?","",fn)
    # 迭代清洗尾部（先去尾横杠，再去 测定标准操作程序 等，反复至稳定）
    for _ in range(4):
        new = re.sub(r"-+$","",fn).strip()
        new = re.sub(r"(?:测定标准操作程序|标准测定程序|标准操作程序|测定)$","",new).strip()
        if new == fn:
            break
        fn = new
    fn = re.sub(r"^(?:酶法|免疫透射比浊法|免疫比浊法|免疫乳胶比浊法|乳胶比浊法|胶乳免疫比浊法|间接离子选择电极法|离子选择电极法|GPO-PAP\s*法|直接法|盒直接法|色原底物法|乳酸底物法|亚铁嗪法|PNP-G7底物法|MPT底物法|NPP\s*底物-?AMP\s*缓冲液法|重氮法|循环酶法|过氧化物酶法|丁酰硫代胆碱底物法|肌氨酸氧化酶法|邻苯三酚红钼法|乳酸氧化酶法|酶偶联物法|速率法|偶氮砷3法|磷钼酸盐法|过氧化氢酶法|溴甲酚绿法|双缩脲法|乳酸脱氢酶法|MDH法|FAPGG底物法|PAPS\s*显色法|脱氢酶法|Nitroso-PSAP)\s*","",fn)
    fn = re.sub(r"^(?:血清|血浆|尿液|脑脊液|全血)","",fn)
    return fn.strip()

_alias = {
    "载脂蛋白A":"载脂蛋白A1",
    "丙氨酸氨基转氨酶":"丙氨酸氨基转移酶",
    "天门冬氨酸氨基转氨酶":"天门冬氨酸氨基转移酶",
    "γ-谷氨酰基转肽酶":"γ-谷氨酰基转移酶",
    "α-淀粉酶":"淀粉酶",
    "脂蛋白a测定":"脂蛋白a",
    "钾离子":"钾","钠离子":"钠","氯离子":"氯",
    "磷":"无机磷",
    "脑脊液乳酸":"乳酸",
    "C反应蛋白":"超敏C反应蛋白",
    "不饱和铁结合力":"不饱和铁结合力/总铁结合力",
    "β羟丁酸":"β-羟丁酸",
    "尿液轻链κ":"κ轻链（血清）","尿液轻链λ":"λ轻链（血清）",
    "免疫球蛋白IgG4":"免疫球蛋白G4",
    "尿或脑脊液总蛋白":"总蛋白",
    "淀粉样蛋白A":"血清淀粉样蛋白A",
    "锌测定":"锌",
}
# 特殊：系统无此规范项的，按仪器档案变体归属指定
FORCE = {
    "腺苷脱氨酶":{"AU5800"},
    # 用户更正：尿酸在 AU5800急诊(1005) 也做 -> 三台联动
    "尿酸":{"AU58-1","AU58-2","AU5800"},
}
# 特殊：明确归属非 AU58 系列的机器（用户指定编号，系统仪器管理未登记）
SPECIAL = {
    "N-乙酰-β-D-氨基葡萄糖苷酶": {"sop":"1013", "name":"日立HT7600",
        "title":"日立HT7600全自动生化分析仪操作作业指导书"},
}

def resolve(proj):
    if proj in FORCE: return FORCE[proj]
    name = _alias.get(proj, proj)
    return proj_to_machines.get(name)

MACHINE_INFO = {fam:{"sop":sop,"name":name} for iid,(fam,sop,name) in INSTS.items()}

def build_ref(machines, pre="MHZYY-JYK-SM-"):
    ordered = sorted(machines, key=lambda m: int(MACHINE_INFO[m]["sop"]))
    sops = "/".join(MACHINE_INFO[m]["sop"] for m in ordered)
    names = "/".join(MACHINE_INFO[m]["name"] for m in ordered)
    return f"{pre}SOP-{sops}《贝克曼{names}全自动生化分析仪操作作业指导书》"

# ---------- docx run 保留替换 ----------
import copy as _copy
def copy_run_style(dst, src):
    f = src.font
    df = dst.font
    if f.name: df.name = f.name
    if f.size: df.size = f.size
    if f.bold is not None: df.bold = f.bold
    if f.italic is not None: df.italic = f.italic
    if f.underline is not None: df.underline = f.underline
    if f.color and f.color.rgb is not None:
        df.color.rgb = f.color.rgb
    # 中文字体：深拷贝 rPr（必须 copy，否则会移动源 run 的 rPr 导致后续复用丢失格式）
    rpr = src._element.find(qn('w:rPr'))
    if rpr is not None:
        dst._element.append(_copy.deepcopy(rpr))

PATTERN = re.compile(r"(?P<pre>(?:MHZYY|HZYY)-JYK-SM-)?SOP-1002(?P<de>的)?《?AU5800生化分析仪标准操作程序》?")
# 泛称仪器 SOP 引用（用于 NA 文件）：匹配「仪器(标准)?操作规程」或「仪器(标准)?操作程序」子串，保留前面的 详见/参照
GEN_PATTERN = re.compile(r"仪器(标准)?(?:操作程序|操作规程)")

def replace_paragraph(par, pattern, make_repl):
    runs = par.runs
    if not runs: return False
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches: return False
    # char -> run index
    char_to_run = []
    for ri,r in enumerate(runs):
        char_to_run.extend([ri]*len(r.text))
    orig_runs = list(runs)
    segments = []  # (text, style_run_index)
    last = 0
    for m in matches:
        s,e = m.span()
        # text before
        i = last
        while i < s:
            ri = char_to_run[i]
            seg_start = i
            while i < s and char_to_run[i]==ri: i+=1
            segments.append((full[seg_start:i], ri))
        # replacement (style = run at s)
        rs = char_to_run[s]
        newtext = make_repl(m)
        segments.append((newtext, rs))
        last = e
    # trailing
    i = last
    while i < len(full):
        ri = char_to_run[i]
        seg_start = i
        while i < len(full) and char_to_run[i]==ri: i+=1
        segments.append((full[seg_start:i], ri))
    # apply
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    for text, ri in segments:
        if text=="": continue
        nr = par.add_run(text)
        copy_run_style(nr, orig_runs[ri])
    return True

# ---------- main ----------
files = sorted([f for f in os.listdir(FOLDER) if f.lower().endswith(".docx") and not f.startswith("~$")])
ONLY = os.environ.get("ONLY")
if ONLY: files = [f for f in files if ONLY in f]
APPLY = os.environ.get("APPLY")=="1"
os.makedirs(BACKUP, exist_ok=True)

report = []
for fn in files:
    proj = parse_project(fn)
    path = os.path.join(FOLDER, fn)
    doc = Document(path)
    if proj in SPECIAL:
        sp = SPECIAL[proj]
        sops = sp["sop"]; title = sp["title"]; machines = [sp["name"]]
    else:
        machines = resolve(proj)
        if machines is None:
            report.append({"file":fn,"project":proj,"machines":None,"refs":0,"action":"SKIP(无法匹配机器)"})
            continue
        ordered = sorted(machines, key=lambda m: int(MACHINE_INFO[m]["sop"]))
        sops = "/".join(MACHINE_INFO[m]["sop"] for m in ordered)
        title = "贝克曼" + "/".join(MACHINE_INFO[m]["name"] for m in ordered) + "全自动生化分析仪操作作业指导书"
    # 两种替换目标
    def sop_repl(m):
        pre = m.group("pre") or "MHZYY-JYK-SM-"
        de = m.group("de") or ""
        return f"{pre}SOP-{sops}{de}《{title}》"
    def gen_repl(m):
        return f"MHZYY-JYK-SM-SOP-{sops}《{title}》"
    changed = 0
    for p in doc.paragraphs:
        if replace_paragraph(p, PATTERN, sop_repl): changed += 1
        if replace_paragraph(p, GEN_PATTERN, gen_repl): changed += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_paragraph(p, PATTERN, sop_repl): changed += 1
                    if replace_paragraph(p, GEN_PATTERN, gen_repl): changed += 1
    if changed == 0:
        report.append({"file":fn,"project":proj,"machines":sorted(machines),"refs":0,"action":"SKIP(无任何仪器SOP引用)"})
        continue
    if APPLY:
        shutil.copy2(path, os.path.join(BACKUP, fn))
        doc.save(path)
    report.append({"file":fn,"project":proj,"machines":sorted(machines),"refs":changed,
                   "new_ref":f"SOP-{sops}《{title}》","changed_paras":changed,
                   "action":"APPLY" if APPLY else "DRYRUN"})

print(f"{'APPLY' if APPLY else 'DRY-RUN'}  files={len(report)}")
print(f"{'FILE':52} {'PROJECT':18} {'MACHINES':18} {'REFS':4} ACTION")
for r in report:
    print(f"{r['file'][:50]:52} {str(r['project'])[:16]:18} {str(r['machines'])[:16]:18} {r['refs']:<4} {r['action']}")
# summary
skip_nomachine = [r for r in report if r['action'].startswith('SKIP(无法')]
skip_noref = [r for r in report if r['action']=='SKIP(无任何仪器SOP引用)']
done = [r for r in report if r['action'] in ('APPLY','DRYRUN')]
print(f"\n将替换(含dry): {len(done)}  无任何仪器SOP引用跳过: {len(skip_noref)}  无法匹配机器跳过: {len(skip_nomachine)}")
for r in skip_nomachine: print("  无法匹配:", r['file'], "| project:", r['project'])
print("\n=== 跳过(无任何仪器SOP引用)的文件 ===")
for r in skip_noref:
    print(f"  {r['file']}")
print("\n=== 将替换文件的引用目标 ===")
for r in done:
    print(f"  {r['file'][:46]:48} -> {r['new_ref']}")
