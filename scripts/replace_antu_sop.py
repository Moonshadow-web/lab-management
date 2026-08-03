#!/usr/bin/env python
# 安图(Autobio)仪器升级：A2000 / A2000PLUS -> A6200
# 规则（与用户确认）：
#  ① 标题/文档内无空格形式：AutolumoA2000 -> AutoLumoA6200  (改数字+大写L修正笔误)
#  ② 正文带空格+PLUS形式：AutoLumo A2000PLUS -> AutoLumo A6200
#  ③ 仪器SOP书名引用：
#     (详见编号)?(MHZYY|HZYY)-JYK-SM-SOP-1019《安图A2000PLUS免疫分析仪标准操作作业指导书》
#       -> 详见编号MHZYY-JYK-SM-SOP-2010《安图A6200免疫分析仪操作作业指导书》
#        （统一加"详见编号"前缀、书名去"标准"、前缀按用户示例改 MHZYY、1019->2010）
# 仅替换上述精确子串；run 级 rPr 深拷贝保留格式。APPLY 前备份原文件到 安图_BACKUP/。
# ONLY=文件名片段 可单文件处理。
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn
import copy as _copy

FOLDER = r"C:/Users/81526/Desktop/待办/安图"
BACKUP = r"C:/Users/81526/Desktop/待办/安图_BACKUP"

# ---------- run 级保留替换 ----------
def copy_run_style(dst, src):
    f = src.font; df = dst.font
    if f.name: df.name = f.name
    if f.size: df.size = f.size
    if f.bold is not None: df.bold = f.bold
    if f.italic is not None: df.italic = f.italic
    if f.underline is not None: df.underline = f.underline
    if f.color and f.color.rgb is not None: df.color.rgb = f.color.rgb
    rpr = src._element.find(qn('w:rPr'))
    if rpr is not None: dst._element.append(_copy.deepcopy(rpr))

def replace_paragraph(par, pattern, make_repl):
    runs = par.runs
    if not runs: return False
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches: return False
    char_to_run = []
    for ri, r in enumerate(runs):
        char_to_run.extend([ri] * len(r.text))
    orig_runs = list(runs)
    segments = []; last = 0
    for m in matches:
        s, e = m.span()
        i = last
        while i < s:
            ri = char_to_run[i]; seg_start = i
            while i < s and char_to_run[i] == ri: i += 1
            segments.append((full[seg_start:i], ri))
        rs = char_to_run[s]
        segments.append((make_repl(m), rs))
        last = e
    i = last
    while i < len(full):
        ri = char_to_run[i]; seg_start = i
        while i < len(full) and char_to_run[i] == ri: i += 1
        segments.append((full[seg_start:i], ri))
    for r in list(par.runs): r._element.getparent().remove(r._element)
    for text, ri in segments:
        if text == "": continue
        nr = par.add_run(text); copy_run_style(nr, orig_runs[ri])
    return True

# ---------- 替换规则 ----------
# ① 无空格形式（标题/文档内）：AutolumoA2000 -> AutoLumoA6200（忽略大小写，输出大写L）
RULE1 = (re.compile(r"(?i)autolumoA2000"), "AutoLumoA6200")
# ② 带空格+PLUS形式：AutoLumo A2000PLUS -> AutoLumo A6200（精确，大写L）
RULE2 = (re.compile(r"AutoLumo\s*A2000\s*PLUS"), "AutoLumo A6200")
# ③ 仪器SOP书名引用
PAT_SOP = re.compile(
    r"(详见编号)?(?:MHZYY|HZYY)-JYK-SM-SOP-1019《安图A2000PLUS免疫分析仪标准操作作业指导书》"
)
def sop_repl(m):
    return "详见编号MHZYY-JYK-SM-SOP-2010《安图A6200免疫分析仪操作作业指导书》"

RULES = [RULE1, RULE2, (PAT_SOP, sop_repl)]

def process(paras):
    c = 0
    for p in paras:
        for pat, new in RULES:
            if replace_paragraph(p, pat, (lambda m, n=new: n) if isinstance(new, str) else new):
                c += 1
    return c

ONLY = os.environ.get("ONLY")
APPLY = os.environ.get("APPLY") == "1"
os.makedirs(BACKUP, exist_ok=True)
files = sorted(f for f in os.listdir(FOLDER) if f.lower().endswith(".docx") and not f.startswith("~$"))
if ONLY: files = [f for f in files if ONLY in f]

report = []
for fn in files:
    path = os.path.join(FOLDER, fn)
    doc = Document(path)
    changed = 0
    changed += process(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                changed += process(cell.paragraphs)
    if changed == 0:
        report.append((fn, 0, "SKIP(无匹配)"))
        continue
    if APPLY:
        shutil.copy2(path, os.path.join(BACKUP, fn))
        doc.save(path)
    report.append((fn, changed, "APPLY" if APPLY else "DRYRUN"))

print(f"{'APPLY' if APPLY else 'DRY-RUN'}  files={len(report)}")
tot = sum(r[1] for r in report)
print(f"总替换段落/单元格处={tot}")
for fn, c, a in report:
    print(f"  {a:7} {c:3}  {fn[:54]}")
