#!/usr/bin/env python
# 替换 DXI 系列项目 SOP 文档里对通用/旧仪器 SOP 的引用，改为该项目实际使用的 DXI 机器。
import os, re, json, shutil, urllib.request, urllib.parse
from docx import Document
from docx.oxml.ns import qn

FOLDER = r"C:/Users/81526/Desktop/待办/DXI"
BACKUP = r"C:/Users/81526/Desktop/待办/DXI_BACKUP"
BASE="http://lab-management-282724-9-1408547492.sh.run.tcloudbase.com/api/v1"

# ---------- API ----------
def req(method,path,data=None,token=""):
    url=BASE+path; h={"Accept":"application/json"}
    if token: h["Authorization"]=f"Bearer {token}"
    body=None
    if data is not None:
        if method=="POST" and path=="/auth/login":
            body=urllib.parse.urlencode(data).encode(); h["Content-Type"]="application/x-www-form-urlencoded"
        else:
            body=json.dumps(data,ensure_ascii=False).encode(); h["Content-Type"]="application/json"
    r=urllib.request.Request(url,data=body,headers=h,method=method)
    try:
        resp=urllib.request.urlopen(r,timeout=120); return resp.status,json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        return e.code,e.read().decode()[:300]
st,l=req("POST","/auth/login",{"username":"jinzizheng","password":"Jzz6827556"}); tok=l["access_token"]

# 主表 id -> 规范名（仪器档案接口的 name 字段部分缺失，必须用 id 反查主表）
st,r=req("GET","/test-items?page=1&page_size=5000",token=tok)
id_to_name={}
for it in r.get("items",[]):
    id_to_name[it["id"]]=(it.get("name") or "").strip()

# DXI 机器： id -> (档案名, SOP尾号, 文档显示名)
INSTS = {
    69:("1号机","2004","DXI800-1"),
    70:("2号机","2005","DXI800-2"),
    71:("3号机","2006","DXI800-3"),
    72:("4号机","2007","DXI800-4"),
    73:("急诊","2008","DXI800急"),
    74:("唐筛","2009","DXI800唐"),
}
proj_to_machines = {}
for iid,(fam,sop,name) in INSTS.items():
    st,r2=req("GET",f"/instruments/{iid}/test-items",token=tok)
    for it in r2:
        nm=id_to_name.get(it.get("id"))
        if nm: proj_to_machines.setdefault(nm,set()).add(fam)

# ---------- 文件名 -> 项目核心名 ----------
def parse_project(filename):
    fn = filename.rsplit(".",1)[0]
    fn = re.sub(r"\s+NA$","",fn).strip()
    fn = re.sub(r"^SM-SOP-\d+[\s\-]*","",fn)
    fn = re.sub(r"^贝克曼DXI800检测系统","",fn)
    # 方法前缀（可能没有，如 524 直接以样品类型开头）
    fn = re.sub(r"^(?:化学发光法|电化学发光法|IGRA法|免疫分析|酶联免疫吸附法)\s*","",fn)
    # 连续多个样品类型前缀（如 血清血浆）全部去掉
    fn = re.sub(r"^(?:血清|血浆|尿液|全血|胸腹水|脑脊液)+","",fn)
    # 尾部 kit/方法后缀
    fn = re.sub(r"\s+沃文特$","",fn).strip()
    fn = re.sub(r"FLOW$","",fn).strip()
    # 去尾
    for _ in range(4):
        new = re.sub(r"-+$","",fn).strip()
        new = re.sub(r"(?:测定标准操作程序|标准测定程序|标准操作程序|操作程序|测定)$","",new).strip()
        if new == fn: break
        fn = new
    return fn.strip()

_alias = {
    "总β-人绒毛膜促性腺激素β-HCG":"β人绒毛膜促性腺激素",
    "β-HCG":"β人绒毛膜促性腺激素",
    "雌二醇E2":"雌二醇",
    "未结合雌三醇u-E3":"非结合型雌三醇",
    "未结合雌三醇E3":"非结合型雌三醇",
    "总25-羟维生素D":"维生素D",
    "25-羟维生素D":"维生素D",
    "全段甲状旁腺激素iPTH":"甲状旁腺激素",
    "甲状旁腺激素iPTH":"甲状旁腺激素",
    "超敏肌钙蛋白I":"超敏肌钙蛋白Ⅰ",
    "超敏肌钙蛋白":"超敏肌钙蛋白Ⅰ",
    "可溶性转铁蛋白受体sTfR":"可溶性转铁蛋白受体",
    "sTfR":"可溶性转铁蛋白受体",
    "肌酸激酶同工酶CK-MB":"肌酸激酶同工酶",
    "CK-MB":"肌酸激酶同工酶",
    "甲胎蛋白AFP":"甲胎蛋白",
    "AFP":"甲胎蛋白",
    "总前列腺特异性抗原TPSA":"总前列腺特异性抗原",
    "TPSA":"总前列腺特异性抗原",
    "游离前列腺特异性抗原TPSA":"游离前列腺特异性抗原",
    "游离前列腺特异性抗原fPSA":"游离前列腺特异性抗原",
    "抗缪勒管激素AMH":"抗缪勒管激素",
    "AMH":"抗缪勒管激素",
    "甲状腺球蛋白抗体TgAb":"抗甲状腺球蛋白抗体",
    "TgAb":"抗甲状腺球蛋白抗体",
    "抗甲状腺过氧化物酶抗体TPOAb":"抗甲状腺过氧化物酶抗体",
    "TPOAb":"抗甲状腺过氧化物酶抗体",
    "促红细胞生成素EPO":"促红细胞生成素",
    "EPO":"促红细胞生成素",
    "B型钠尿肽BNP":"B型钠尿肽",
    "BNP":"B型钠尿肽",
    "肌红蛋白MYO":"肌红蛋白",
    "MYO":"肌红蛋白",
    "叶酸FLOW":"叶酸",
    "内因子抗体IFAb":"抗内因子抗体",
    "IFAb":"抗内因子抗体",
    "总甲状腺素TT4":"甲状腺素",
    "总三碘甲状腺原氨酸TT3":"三碘甲状腺原氨酸",
    "游离甲状腺素FT4":"游离甲状腺素",
    "游离三碘甲状腺原氨酸FT3":"游离三碘甲状腺原氨酸",
    "促甲状腺激素TSH":"促甲状腺激素",
    "白介素6":"白介素-6",
    # 安全兜底：若样品类型未剥干净
    "血浆降钙素原":"降钙素原",
    "血清降钙素原":"降钙素原",
    "血清血浆降钙素原":"降钙素原",
    "血浆白介素6":"白介素-6",
    "血清白介素6":"白介素-6",
    "血清血浆白介素6":"白介素-6",
    "血浆铁蛋白":"铁蛋白",
    "血清铁蛋白":"铁蛋白",
    "血清血浆铁蛋白":"铁蛋白",
    "血浆可溶性转铁蛋白受体":"可溶性转铁蛋白受体",
    "血浆内因子抗体":"抗内因子抗体",
    "血浆叶酸":"叶酸",
    "血清血浆叶酸":"叶酸",
    "血浆维生素B12":"维生素B12",
    "血清血浆维生素B12":"维生素B12",
    "血浆B型钠尿肽":"B型钠尿肽",
    "血清血浆B型钠尿肽":"B型钠尿肽",
    "血浆肌红蛋白":"肌红蛋白",
    "血清血浆肌红蛋白":"肌红蛋白",
    "血浆超敏肌钙蛋白I":"超敏肌钙蛋白Ⅰ",
    "血清血浆超敏肌钙蛋白I":"超敏肌钙蛋白Ⅰ",
    "血浆超敏肌钙蛋白":"超敏肌钙蛋白Ⅰ",
    "血清血浆超敏肌钙蛋白":"超敏肌钙蛋白Ⅰ",
    "血浆肌酸激酶同工酶":"肌酸激酶同工酶",
    "血清血浆肌酸激酶同工酶":"肌酸激酶同工酶",
    "血清降钙素":"降钙素",
    "血浆降钙素":"降钙素",
}

def resolve(proj):
    name = _alias.get(proj, proj)
    return proj_to_machines.get(name)

MACHINE_INFO = {fam:{"sop":sop,"name":name} for iid,(fam,sop,name) in INSTS.items()}

def build_ref(machines, pre="MHZYY-JYK-SM-"):
    ordered = sorted(machines, key=lambda m: int(MACHINE_INFO[m]["sop"]))
    sops = "/".join(MACHINE_INFO[m]["sop"] for m in ordered)
    names = "/".join(MACHINE_INFO[m]["name"] for m in ordered)
    return f"{pre}SOP-{sops}《贝克曼{names}化学发光仪标准操作作业指导书》"

# ---------- docx run 保留替换 ----------
import copy as _copy
def copy_run_style(dst, src):
    f = src.font
    df = dst.font
    if f.name: df.name = f.name
    if f.size: df.size = f.size
    if f.bold is not None: df.bold = f.bold
    if f.italic is not None: df.italic = f.italic
    if f.underline is not None: df.underline = f.underline
    if f.color and f.color.rgb is not None:
        df.color.rgb = f.color.rgb
    rpr = src._element.find(qn('w:rPr'))
    if rpr is not None:
        dst._element.append(_copy.deepcopy(rpr))

# 匹配显式仪器 SOP 引用：SOP-1004、SOP-1006、SOP-1004/1006 等，以及带/不带前缀和「的」
PATTERN = re.compile(
    r"(?P<pre>(?:MHZYY|HZYY)-JYK-SM-)?"
    r"SOP-(?P<sops>\d+(?:/\d+)*)"
    r"(?P<de>的)?"
    r"《?贝克曼DXI800[^《》]*化学发光仪标准操作作业指导书》?"
)
# 泛称仪器 SOP 引用（可能出现在 NA 或旧段落中）
# 注意：直接替换会误伤「化学发光免疫分析仪标准操作程序」之类标题，
# 默认关闭；仅当用户显式 GEN=1 才启用，且需人工复核。
GEN_PATTERN = re.compile(r"仪器(?:标准)?(?:操作程序|操作规程)")

def replace_paragraph(par, pattern, make_repl):
    runs = par.runs
    if not runs: return False
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches: return False
    char_to_run = []
    for ri,r in enumerate(runs):
        char_to_run.extend([ri]*len(r.text))
    orig_runs = list(runs)
    segments = []
    last = 0
    for m in matches:
        s,e = m.span()
        i = last
        while i < s:
            ri = char_to_run[i]
            seg_start = i
            while i < s and char_to_run[i]==ri: i+=1
            segments.append((full[seg_start:i], ri))
        rs = char_to_run[s]
        newtext = make_repl(m)
        segments.append((newtext, rs))
        last = e
    i = last
    while i < len(full):
        ri = char_to_run[i]
        seg_start = i
        while i < len(full) and char_to_run[i]==ri: i+=1
        segments.append((full[seg_start:i], ri))
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    for text, ri in segments:
        if text=="": continue
        nr = par.add_run(text)
        copy_run_style(nr, orig_runs[ri])
    return True

# ---------- main ----------
files = sorted([f for f in os.listdir(FOLDER) if f.lower().endswith(".docx") and not f.startswith("~$")])
ONLY = os.environ.get("ONLY")
if ONLY: files = [f for f in files if ONLY in f]
APPLY = os.environ.get("APPLY")=="1"
USE_GEN = os.environ.get("GEN")=="1"   # 默认关闭裸「仪器操作规程」替换，避免误伤标题
os.makedirs(BACKUP, exist_ok=True)

report = []
for fn in files:
    proj = parse_project(fn)
    path = os.path.join(FOLDER, fn)
    doc = Document(path)
    machines = resolve(proj)
    if machines is None:
        report.append({"file":fn,"project":proj,"machines":None,"refs":0,"action":"SKIP(无法匹配机器)"})
        continue
    ordered = sorted(machines, key=lambda m: int(MACHINE_INFO[m]["sop"]))
    sops = "/".join(MACHINE_INFO[m]["sop"] for m in ordered)
    title = "贝克曼" + "/".join(MACHINE_INFO[m]["name"] for m in ordered) + "化学发光仪标准操作作业指导书"
    def sop_repl(m):
        pre = m.group("pre") or "MHZYY-JYK-SM-"
        de = m.group("de") or ""
        return f"{pre}SOP-{sops}{de}《{title}》"
    def gen_repl(m):
        return f"MHZYY-JYK-SM-SOP-{sops}《{title}》"
    changed = 0
    for p in doc.paragraphs:
        if replace_paragraph(p, PATTERN, sop_repl): changed += 1
        if USE_GEN and replace_paragraph(p, GEN_PATTERN, gen_repl): changed += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_paragraph(p, PATTERN, sop_repl): changed += 1
                    if USE_GEN and replace_paragraph(p, GEN_PATTERN, gen_repl): changed += 1
    if changed == 0:
        report.append({"file":fn,"project":proj,"machines":sorted(machines),"refs":0,"action":"SKIP(无任何仪器SOP引用)"})
        continue
    if APPLY:
        shutil.copy2(path, os.path.join(BACKUP, fn))
        doc.save(path)
    report.append({"file":fn,"project":proj,"machines":sorted(machines),"refs":changed,
                   "new_ref":f"SOP-{sops}《{title}》","changed_paras":changed,
                   "action":"APPLY" if APPLY else "DRYRUN"})

print(f"{'APPLY' if APPLY else 'DRY-RUN'}  files={len(report)}")
print(f"{'FILE':52} {'PROJECT':22} {'MACHINES':24} {'REFS':4} ACTION")
for r in report:
    print(f"{r['file'][:50]:52} {str(r['project'])[:20]:22} {str(r['machines'])[:22]:24} {r['refs']:<4} {r['action']}")
skip_nomachine = [r for r in report if r['action'].startswith('SKIP(无法')]
skip_noref = [r for r in report if r['action']=='SKIP(无任何仪器SOP引用)']
done = [r for r in report if r['action'] in ('APPLY','DRYRUN')]
print(f"\n将替换(含dry): {len(done)}  无任何仪器SOP引用跳过: {len(skip_noref)}  无法匹配机器跳过: {len(skip_nomachine)}")
if skip_nomachine:
    print("\n=== 无法匹配机器的文件 ===")
    for r in skip_nomachine: print(f"  {r['file']} | project: {r['project']}")
if skip_noref:
    print("\n=== 无任何仪器SOP引用的文件 ===")
    for r in skip_noref: print(f"  {r['file']}")
print("\n=== 将替换文件的引用目标 ===")
for r in done:
    print(f"  {r['file'][:46]:48} -> {r['new_ref']}")
