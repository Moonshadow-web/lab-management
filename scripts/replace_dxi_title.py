#!/usr/bin/env python
# 把 DXI 文档里标题/章节名中的「测定标准操作程序」改为「测定操作作业指导书」（用户要求），
# 不动仪器SOP引用（已是「标准操作作业指导书」）。
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn
import copy as _copy

FOLDER = r"C:/Users/81526/Desktop/待办/DXI"
BACKUP = r"C:/Users/81526/Desktop/待办/DXI_BACKUP_TITLE"
OLD = "测定标准操作程序"
NEW = "测定操作作业指导书"
PATTERN = re.compile(re.escape(OLD))

def copy_run_style(dst, src):
    f = src.font; df = dst.font
    if f.name: df.name = f.name
    if f.size: df.size = f.size
    if f.bold is not None: df.bold = f.bold
    if f.italic is not None: df.italic = f.italic
    if f.underline is not None: df.underline = f.underline
    if f.color and f.color.rgb is not None: df.color.rgb = f.color.rgb
    rpr = src._element.find(qn('w:rPr'))
    if rpr is not None: dst._element.append(_copy.deepcopy(rpr))

def replace_paragraph(par, pattern, make_repl):
    runs = par.runs
    if not runs: return False
    full = "".join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches: return False
    char_to_run = []
    for ri, r in enumerate(runs):
        char_to_run.extend([ri] * len(r.text))
    orig_runs = list(runs)
    segments = []; last = 0
    for m in matches:
        s, e = m.span()
        i = last
        while i < s:
            ri = char_to_run[i]; seg_start = i
            while i < s and char_to_run[i] == ri: i += 1
            segments.append((full[seg_start:i], ri))
        segments.append((make_repl(m), char_to_run[s]))
        last = e
    i = last
    while i < len(full):
        ri = char_to_run[i]; seg_start = i
        while i < len(full) and char_to_run[i] == ri: i += 1
        segments.append((full[seg_start:i], ri))
    for r in list(par.runs): r._element.getparent().remove(r._element)
    for text, ri in segments:
        if text == "": continue
        nr = par.add_run(text); copy_run_style(nr, orig_runs[ri])
    return True

def process(paras):
    c = 0
    for p in paras:
        if replace_paragraph(p, PATTERN, lambda m: NEW): c += 1
    return c

def header_footer_match_count(doc):
    n = 0
    for rel in doc.part.rels.values():
        if "header" in rel.reltype or "footer" in rel.reltype:
            try:
                part = rel.target_part
                blob = part.blob
                n += blob.decode("utf-8", "ignore").count(OLD)
            except Exception:
                pass
    return n

ONLY = os.environ.get("ONLY")
APPLY = os.environ.get("APPLY") == "1"
os.makedirs(BACKUP, exist_ok=True)
files = sorted(f for f in os.listdir(FOLDER) if f.lower().endswith(".docx") and not f.startswith("~$"))
if ONLY: files = [f for f in files if ONLY in f]
report = []
for fn in files:
    path = os.path.join(FOLDER, fn)
    doc = Document(path)
    changed = 0
    changed += process(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                changed += process(cell.paragraphs)
    hf = header_footer_match_count(doc)
    if changed == 0 and hf == 0:
        report.append((fn, 0, 0, "SKIP"))
        continue
    if APPLY:
        shutil.copy2(path, os.path.join(BACKUP, fn))
        doc.save(path)
    report.append((fn, changed, hf, "APPLY" if APPLY else "DRYRUN"))
print(f"{'APPLY' if APPLY else 'DRY-RUN'} files={len(report)}")
total_body = sum(r[1] for r in report)
total_hf = sum(r[2] for r in report)
print(f"body替换处={total_body}  页眉/页脚残留={total_hf}")
for fn, c, hf, a in report:
    print(f"  {a:7} body={c:3} hf={hf:2}  {fn[:54]}")
