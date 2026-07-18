"""诊断：打印 3 份 DOCX 的实际表格结构（含合并单元格 gridSpan/vMerge）。"""
import docx
from docx.oxml.ns import qn

FILES = {
    "IFAb": r"C:/Users/81526/Desktop/2026年6月室内室间比对结果(2)/2026年6月室内室间比对结果/BG-SM-CZ-019-定量项目室间比对结果记录及分析报告表IFAb.docx",
    "p2PSA": r"C:/Users/81526/Desktop/2026年6月室内室间比对结果(2)/2026年6月室内室间比对结果/BG-SM-CZ-019-定量项目室间比对结果记录及分析报告表p2PSA.docx",
    "EPO": r"C:/Users/81526/Desktop/2026年6月室内室间比对结果(2)/2026年6月室内室间比对结果/BG-SM-CZ-019-定量项目室间比对结果记录及分析报告表EPO.docx",
}


def cell_spans(tc):
    gs = tc.find(qn('w:tcPr') + '/' + qn('w:gridSpan'))
    span = int(gs.get(qn('w:val'))) if gs is not None else 1
    vm = tc.find(qn('w:tcPr') + '/' + qn('w:vMerge'))
    vmerge = vm.get(qn('w:val')) if vm is not None else None  # 'restart' / 'continue' / None
    return span, vmerge


def dump(path):
    d = docx.Document(path)
    print("=" * 80)
    print("TABLES:", len(d.tables))
    for ti, t in enumerate(d.tables):
        print(f"\n--- TABLE {ti}: rows={len(t.rows)} cols(inferred)={len(t.columns)} ---")
        for ri, r in enumerate(t.rows):
            cells = r.cells
            parts = []
            for ci, c in enumerate(cells):
                # 用底层 tc 判断 span（cells 会重复合并单元格，无法直接看 span）
                tc = c._tc
                span, vmerge = cell_spans(tc)
                txt = (c.text or "").strip().replace("\n", "\\n")
                tag = ""
                if span > 1:
                    tag = f"[gs{span}]"
                if vmerge:
                    tag = f"[vm:{vmerge}]"
                parts.append(f"c{ci}:{txt!r}{tag}")
            print(f"  r{ri}: " + " | ".join(parts))


for k, v in FILES.items():
    print("\n\n########## FILE:", k, "##########")
    dump(v)
